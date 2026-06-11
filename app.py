#!/usr/bin/env python
"""HR System Flask Application"""
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from werkzeug.security import check_password_hash, generate_password_hash
from datetime import datetime, timezone, timedelta
from pathlib import Path
from email.message import EmailMessage
import base64
import csv
import io
import json
import os
import re
import secrets
import smtplib
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

app = Flask(__name__)
BRANDING_DIR = Path(__file__).parent / "uploads" / "branding"
BRANDING_DIR.mkdir(parents=True, exist_ok=True)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret-key-12345")
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
# Session cookie expires when browser/tab closes (no permanent session flag on login)

DATA_DIR = Path(__file__).parent / "data"
DATA_DIR.mkdir(exist_ok=True)

# In-memory storage
USERS = {}
JOBS = {}
APPLICATIONS = {}
ACTIVITY_LOGS = []
EMAIL_HISTORY = []
SETTINGS = {}
INTERVIEWER_PANELS = {}

HR_ROLES = ("hr", "hr_staff")


def _generate_panel_access_code():
    """One-time access code shared by all interviewers on a panel link."""
    return f"{secrets.randbelow(90000000) + 10000000}"

DEFAULT_SETTINGS = {
    "company_name": "Grain Marketing Board",
    "support_email": "hr@company.com",
    "default_view": "dashboard",
    "items_per_page": 25,
    "notification_email_enabled": True,
    "application_updates_enabled": True,
    "interview_reminders_enabled": True,
    "weekly_reports_enabled": False,
    "mail_server": os.environ.get("MAIL_SERVER", "smtp.gmail.com"),
    "mail_port": int(os.environ.get("MAIL_PORT", 587)),
    "mail_use_tls": str(os.environ.get("MAIL_USE_TLS", "true")).lower() in ("1", "true", "yes", "on"),
    "mail_username": os.environ.get("MAIL_USERNAME", ""),
    "mail_password": os.environ.get("MAIL_PASSWORD", ""),
    "mail_default_sender": os.environ.get("MAIL_DEFAULT_SENDER", os.environ.get("MAIL_USERNAME", "noreply@company.com")),
    "interview_default_message": (
        "Dear {candidate_name},\n\n"
        "You are invited to an interview for the {job_title} role on {interview_date} at {interview_time}.\n"
        "Interview type: {interview_type}.\n"
        "Location/Link: {interview_location}.\n\n"
        "Regards,\nHR Team"
    ),
}

# ================= DB =================
def connect_db():
    return psycopg2.connect(
        host="localhost",
        database="gate_system",
        user="postgres",
        password="juva"
    )

def _settings_file():
    return DATA_DIR / "settings.json"


def _load_dotenv():
    env_path = Path(__file__).parent / ".env"
    if not env_path.exists():
        return
    for raw_line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def _apply_env_mail_settings():
    env_map = {
        "mail_server": os.environ.get("MAIL_SERVER") or os.environ.get("SMTP_HOST"),
        "mail_port": os.environ.get("MAIL_PORT") or os.environ.get("SMTP_PORT"),
        "mail_username": os.environ.get("MAIL_USERNAME") or os.environ.get("SMTP_USER"),
        "mail_password": os.environ.get("MAIL_PASSWORD") or os.environ.get("SMTP_PASS"),
        "mail_default_sender": os.environ.get("MAIL_DEFAULT_SENDER") or os.environ.get("SMTP_FROM"),
    }
    for key, value in env_map.items():
        if value is None or str(value).strip() == "":
            continue
        if key == "mail_port":
            SETTINGS[key] = int(value)
        else:
            SETTINGS[key] = str(value).strip()
    if str(SETTINGS.get("mail_port")) == "465":
        SETTINGS["mail_use_tls"] = False


_load_dotenv()


def _save_json(path, payload):
    with open(path, "w") as file_obj:
        json.dump(payload, file_obj, indent=2, default=str)


def _to_bool(value):
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in ("1", "true", "yes", "on")


def _normalize_requirements(value):
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        return [item.strip() for item in value.split(",") if item.strip()]
    return []


def _parse_iso_datetime(value):
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=timezone.utc)
    if isinstance(value, str) and value:
        try:
            parsed = datetime.fromisoformat(value)
            return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
        except Exception:
            return None
    return None


def _screening_text_for_job(job):
    requirements = _normalize_requirements(job.get("requirements"))
    custom_criteria = (job.get("screening_criteria") or job.get("custom_screening_criteria") or "").strip()
    parts = [
        job.get("title", ""),
        job.get("description", ""),
        f"Requirements: {', '.join(requirements)}" if requirements else "",
        custom_criteria,
    ]
    return "\n".join(part for part in parts if part)


def _candidate_full_name(app_data):
    """Build display name without duplicating surname when name already includes it."""
    if not app_data:
        return "Candidate"
    name = (app_data.get("name") or "").strip()
    surname = (app_data.get("surname") or "").strip()
    if not surname:
        return name or "Candidate"
    if not name:
        return surname
    name_lower = name.lower()
    surname_lower = surname.lower()
    if name_lower == surname_lower:
        return name
    if name_lower.endswith(surname_lower):
        return name
    name_parts = name_lower.split()
    if surname_lower in name_parts:
        return name
    return f"{name} {surname}".strip()


def _normalize_screening_status(raw_status):
    status = str(raw_status or "pending").strip().lower()
    status_map = {
        "shortlisted": "shortlisted",
        "review": "pending",
        "pending review": "pending",
        "interview": "interview",
        "rejected": "rejected",
    }
    return status_map.get(status, status or "pending")


def _screen_pending_applications(job_ids, trigger="manual"):
    try:
        from services.cv_scoring import ResumeUpload, rank_candidates
    except Exception as exc:
        print(f"Warning: CV scoring service unavailable: {exc}")
        return {"processed": 0, "jobs": []}

    processed = 0
    touched_jobs = []

    for job_id in job_ids:
        job = JOBS.get(job_id)
        if not job:
            continue

        screening_text = _screening_text_for_job(job)
        if not screening_text.strip():
            continue

        job_processed = 0
        for app_data in APPLICATIONS.values():
            if app_data.get("job_id") != job_id:
                continue
            if str(app_data.get("status") or "pending").strip().lower() != "pending":
                continue

            resume_path = Path(app_data.get("resume_path") or "")
            if not resume_path.exists():
                continue

            try:
                result = rank_candidates(
                    screening_text,
                    [ResumeUpload(label=app_data.get("id", "candidate"), path=resume_path)],
                )[0]
            except Exception as exc:
                print(f"Warning: Could not screen application {app_data.get('id')}: {exc}")
                continue

            app_data["score"] = result.score
            app_data["summary"] = result.summary
            app_data["recommendation"] = result.recommendation
            app_data["recommendation_reason"] = result.recommendation_reason
            app_data["matched_skills"] = result.matched_skills
            app_data["missing_skills"] = result.missing_skills
            # Automatically set status based on score: >= 50% = shortlisted, < 50% = rejected
            app_data["status"] = "shortlisted" if result.score >= 50 else "rejected"
            if app_data["status"] == "shortlisted":
                app_data["interview_status"] = "pending_scheduling"
            app_data["screened_at"] = datetime.now(timezone.utc).isoformat()
            app_data["screening_trigger"] = trigger
            job_processed += 1
            processed += 1

        if job_processed:
            touched_jobs.append({"job_id": job_id, "job_title": job.get("title", "Job"), "processed": job_processed})

    if processed:
        save_applications()

    return {"processed": processed, "jobs": touched_jobs}


def _safe_report_filename(requested_filename, default_name, fmt):
    safe_filename = "".join(
        ch for ch in (requested_filename or default_name) if ch.isalnum() or ch in ("-", "_", " ")
    ).strip() or default_name
    extension_map = {"csv": ".csv", "pdf": ".pdf", "docx": ".docx"}
    extension = extension_map.get(fmt, ".csv")
    if not safe_filename.lower().endswith(extension):
        safe_filename = f"{safe_filename}{extension}"
    return safe_filename


def _normalize_report_format(value):
    fmt = str(value or "csv").strip().lower()
    if fmt in {"word", "doc", "docx"}:
        return "docx"
    if fmt == "pdf":
        return "pdf"
    return "csv"


def _build_summary_report_spec(user, status_filter="all", report_title="Applicants Report", report_notes=""):
    filtered_applications = []
    for app_data in APPLICATIONS.values():
        app_status = str(app_data.get("status", "")).strip().lower()
        if status_filter != "all" and app_status != status_filter:
            continue
        filtered_applications.append(app_data)

    rows = []
    for app_data in filtered_applications:
        job = JOBS.get(app_data.get("job_id", ""), {})
        full_name = _candidate_full_name(app_data)
        score = app_data.get("score", 0) or 0
        recommendation = app_data.get("recommendation", "Pending Review") or "Pending Review"
        rows.append([
            app_data.get("id", ""),
            full_name,
            app_data.get("email", ""),
            app_data.get("phone", ""),
            job.get("title", ""),
            f"{score}%",
            (app_data.get("status", "") or "pending").upper(),
            app_data.get("highest_education", ""),
            app_data.get("uploaded_at", ""),
            recommendation,
        ])

    metadata = [
        ["Report Title", report_title],
        ["Generated At", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")],
        ["Generated By", user.get("name", "HR User")],
        ["Status Filter", status_filter.title() if status_filter != "all" else "All Applicants"],
        ["Total Applicants", len(filtered_applications)],
    ]
    if report_notes:
        metadata.append(["Notes", report_notes])

    return {
        "kind": "summary",
        "title": "RECRUITMENT REPORT - APPLICANTS SUMMARY",
        "metadata": metadata,
        "columns": ["ID", "Full Name", "Email", "Phone", "Job Applied", "CV Score %", "Status", "Education", "Applied Date", "Recommendation"],
        "rows": rows,
    }


def _build_cv_analysis_report_spec(user):
    rows = []
    for app_data in APPLICATIONS.values():
        full_name = _candidate_full_name(app_data)
        job = JOBS.get(app_data.get("job_id", ""), {})
        score = app_data.get("score", 0) or 0
        matched_skills = app_data.get("matched_skills", [])
        missing_skills = app_data.get("missing_skills", [])
        education = app_data.get("highest_education", "Not specified")
        summary = app_data.get("summary", "")
        
        # Professional recommendations based on score
        if score >= 70:
            recommendation = "INTERVIEW CANDIDATE"
        elif score >= 50:
            recommendation = "POTENTIAL MATCH"
        else:
            recommendation = "FURTHER REVIEW"

        rows.append([
            full_name,
            app_data.get("email", ""),
            app_data.get("phone", ""),
            job.get("title", "Unknown Position"),
            education,
            f"{score}%",
            (app_data.get("status", "pending")).upper(),
            ", ".join(matched_skills) if matched_skills else "-",
            ", ".join(missing_skills) if missing_skills else "-",
            recommendation,
            str(app_data.get("screened_at", "Not yet screened")).split("T")[0] if app_data.get("screened_at") != "Not yet screened" else "Not yet screened",
        ])

    return {
        "kind": "analysis",
        "title": "RECRUITMENT REPORT - CV ANALYSIS",
        "metadata": [
            ["Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")],
            ["Generated By", user.get("name", "HR User")],
            ["Total Candidates", len(APPLICATIONS)],
        ],
        "columns": ["Candidate Name", "Email", "Phone", "Job Applied", "Education", "CV Score %", "Status", "Matched Skills", "Missing Skills", "Recommendation", "Screened Date"],
        "rows": rows,
    }


def _build_report_csv(spec):
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([spec["title"]])
    for label, value in spec["metadata"]:
        writer.writerow([label, value])
    writer.writerow([])
    writer.writerow(spec["columns"])
    writer.writerows(spec["rows"])
    return output.getvalue()


def _build_report_docx(spec):
    from docx import Document

    document = Document()
    document.add_heading(spec["title"], level=1)
    for label, value in spec["metadata"]:
        document.add_paragraph(f"{label}: {value}")
    document.add_paragraph("")
    table = document.add_table(rows=1, cols=len(spec["columns"]))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column_name in enumerate(spec["columns"]):
        header_cells[index].text = str(column_name)
    for row_values in spec["rows"]:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_values):
            row_cells[index].text = str(value)
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def _build_report_pdf(spec):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = io.BytesIO()
    document = SimpleDocTemplate(output, pagesize=landscape(letter), leftMargin=24, rightMargin=24, topMargin=24, bottomMargin=24)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph(spec["title"], styles["Title"]),
        Spacer(1, 8),
    ]
    for label, value in spec["metadata"]:
        elements.append(Paragraph(f"<b>{label}:</b> {value}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    table_data = [spec["columns"]] + [[str(value) for value in row] for row in spec["rows"]]
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#124e66")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor("#f7fbfc")]),
    ]))
    elements.append(table)
    document.build(elements)
    output.seek(0)
    return output.getvalue()


def _send_report_download(spec, requested_filename, fmt):
    from flask import Response

    safe_filename = _safe_report_filename(requested_filename, "report", fmt)
    if fmt == "docx":
        payload = _build_report_docx(spec)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "pdf":
        payload = _build_report_pdf(spec)
        mimetype = "application/pdf"
    else:
        payload = _build_report_csv(spec)
        mimetype = "text/csv"
    return Response(payload, mimetype=mimetype, headers={"Content-Disposition": f'attachment; filename="{safe_filename}"'})


def _screen_due_jobs():
    now = datetime.now(timezone.utc)
    due_job_ids = []
    for job_id, job in JOBS.items():
        due_date = _parse_iso_datetime(job.get("due_date"))
        if due_date and due_date <= now:
            due_job_ids.append(job_id)
    return _screen_pending_applications(due_job_ids, trigger="due_date") if due_job_ids else {"processed": 0, "jobs": []}


def save_jobs():
    _save_json(DATA_DIR / "jobs.json", JOBS)


def save_applications():
    _save_json(DATA_DIR / "applications.json", APPLICATIONS)


def save_users():
    _save_json(DATA_DIR / "users.json", USERS)


def save_interviewer_panels():
    _save_json(DATA_DIR / "interviewer_panels.json", INTERVIEWER_PANELS)


def save_settings():
    _save_json(_settings_file(), SETTINGS)


def _env_mail_value(*keys):
    for key in keys:
        value = os.environ.get(key)
        if value and str(value).strip():
            return str(value).strip()
    return ""


def _mail_settings():
    username = (SETTINGS.get("mail_username") or "").strip() or _env_mail_value("MAIL_USERNAME", "SMTP_USER")
    password = (SETTINGS.get("mail_password") or "").strip() or _env_mail_value("MAIL_PASSWORD", "SMTP_PASS")
    sender = (SETTINGS.get("mail_default_sender") or "").strip() or _env_mail_value("MAIL_DEFAULT_SENDER", "SMTP_FROM") or username
    port_raw = SETTINGS.get("mail_port") or _env_mail_value("MAIL_PORT", "SMTP_PORT") or DEFAULT_SETTINGS["mail_port"]
    port = int(port_raw)
    use_tls = _to_bool(SETTINGS.get("mail_use_tls", DEFAULT_SETTINGS["mail_use_tls"]))
    if port == 465:
        use_tls = False
    return {
        "server": (SETTINGS.get("mail_server") or "").strip() or _env_mail_value("MAIL_SERVER", "SMTP_HOST") or DEFAULT_SETTINGS["mail_server"],
        "port": port,
        "use_tls": use_tls,
        "username": username,
        "password": password,
        "sender": sender or DEFAULT_SETTINGS["mail_default_sender"],
    }


def _gmb_logo_data_uri():
    custom_logo = (SETTINGS.get("company_logo_data") or "").strip()
    if custom_logo.startswith("data:image"):
        return custom_logo
    logo_path = Path(__file__).parent / "static" / "gmb-logo.svg"
    if logo_path.exists():
        encoded = base64.b64encode(logo_path.read_bytes()).decode("ascii")
        return f"data:image/svg+xml;base64,{encoded}"
    return ""


def _persist_company_logo_file(data_uri=None, raw_bytes=None, mime="image/png"):
    """Save company logo to disk so emails can embed it via CID (data URIs are blocked by Gmail)."""
    filename = None
    if data_uri and data_uri.startswith("data:image"):
        header, encoded = data_uri.split(",", 1)
        mime = header.split(";")[0].split(":")[1] if ":" in header else mime
        ext_map = {"image/png": "png", "image/jpeg": "jpg", "image/jpg": "jpg", "image/gif": "gif", "image/webp": "webp"}
        ext = ext_map.get(mime, "png")
        raw_bytes = base64.b64decode(encoded)
        filename = f"company_logo.{ext}"
    elif raw_bytes:
        ext_map = {"image/png": "png", "image/jpeg": "jpg", "image/gif": "gif", "image/webp": "webp"}
        ext = ext_map.get(mime, "png")
        filename = f"company_logo.{ext}"
    if not filename or not raw_bytes:
        return None
    save_path = BRANDING_DIR / filename
    save_path.write_bytes(raw_bytes)
    SETTINGS["company_logo_file"] = filename
    return filename


def _get_company_logo_for_email():
    """Return (bytes, mime_subtype) for inline email attachment, or (None, None)."""
    filename = (SETTINGS.get("company_logo_file") or "").strip()
    if filename:
        path = BRANDING_DIR / filename
        if path.exists():
            ext = path.suffix.lower().lstrip(".")
            subtype = "jpeg" if ext in ("jpg", "jpeg") else ext
            if subtype in ("png", "jpeg", "gif", "webp"):
                return path.read_bytes(), subtype
    data_uri = (SETTINGS.get("company_logo_data") or "").strip()
    if data_uri.startswith("data:image"):
        saved = _persist_company_logo_file(data_uri=data_uri)
        if saved:
            return _get_company_logo_for_email()
    return None, None


def _personalize_message(template, values):
    """Replace {candidate_name}, (candidate name), [candidate name], etc."""
    if not template:
        return template
    result = template
    field_patterns = [
        ("candidate_name", values.get("candidate_name") or "Candidate"),
        ("job_title", values.get("job_title") or ""),
        ("application_status", values.get("application_status") or ""),
        ("interview_date", values.get("interview_date") or ""),
        ("interview_time", values.get("interview_time") or ""),
        ("interview_type", values.get("interview_type") or ""),
        ("interview_location", values.get("interview_location") or ""),
        ("start_date", values.get("start_date") or ""),
        ("start_time", values.get("start_time") or ""),
        ("location", values.get("location") or ""),
    ]
    for field, replacement in field_patterns:
        label = field.replace("_", " ")
        result = re.sub(r"\{" + re.escape(field) + r"\}", replacement, result, flags=re.IGNORECASE)
        result = re.sub(r"\(\s*" + re.escape(label) + r"\s*\)", replacement, result, flags=re.IGNORECASE)
        result = re.sub(r"\[\s*" + re.escape(label) + r"\s*\]", replacement, result, flags=re.IGNORECASE)
    return result


def _should_add_dear_greeting(message, candidate_name=None):
    if not message:
        return True
    
    clean_msg = re.sub(r'<[^>]*>', '', message).strip().lower()
    if not clean_msg:
        return True
        
    if clean_msg.startswith(("dear", "hello", "hi", "greetings", "to ", "attention")):
        return False
        
    if candidate_name:
        name_lower = candidate_name.strip().lower()
        if name_lower in clean_msg[:150]:
            return False
            
    return True


def _format_email_content(body_text, candidate_name=None, footer_location=None):
    message = (body_text or "").strip()
    if _should_add_dear_greeting(message, candidate_name):
        lines = [f"Dear {candidate_name or 'Candidate'},", "", message]
    else:
        lines = [message]
    if footer_location:
        lines.extend(["", f"Location: {footer_location}"])
    return "\n".join(line for line in lines if line is not None)


def _build_email_html(body_text, candidate_name=None, footer_location=None, use_cid_logo=False):
    company = SETTINGS.get("company_name", DEFAULT_SETTINGS["company_name"])
    if use_cid_logo:
        logo_html = (
            '<img src="cid:company_logo" alt="Company logo" '
            'style="height:56px;max-width:200px;margin-bottom:12px;border-radius:6px;background:#fff;padding:4px;">'
        )
    else:
        logo_uri = _gmb_logo_data_uri()
        logo_html = (
            f'<img src="{logo_uri}" alt="Company logo" style="height:56px;margin-bottom:12px;border-radius:6px;background:#fff;padding:4px;">'
            if logo_uri
            else f'<div style="font-size:22px;font-weight:700;color:#f39c12;letter-spacing:1px;">{company}</div>'
        )
    name = candidate_name or "Candidate"
    message = (body_text or "").strip()
    if not _should_add_dear_greeting(message, candidate_name):
        greeting_html = ""
        body_html = message.replace("\n", "<br>")
    else:
        greeting_html = f'<p style="margin:0 0 16px;font-size:16px;color:#124e66;"><strong>Dear {name},</strong></p>'
        body_html = message.replace("\n", "<br>")
    footer_html = (
        f'<p style="margin-top:24px;padding-top:16px;border-top:1px solid #eee;color:#124e66;font-weight:600;font-size:14px;">Location: {footer_location}</p>'
        if footer_location
        else ""
    )
    return f"""<!DOCTYPE html>
<html><body style="margin:0;padding:0;font-family:Segoe UI,Arial,sans-serif;background:#f5f7fa;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f5f7fa;padding:24px 0;">
    <tr><td align="center">
      <table width="600" cellpadding="0" cellspacing="0" style="background:#ffffff;border-radius:10px;overflow:hidden;box-shadow:0 4px 12px rgba(0,0,0,0.08);">
        <tr><td style="background:linear-gradient(135deg,#124e66,#0a2338);padding:24px;text-align:center;">
          {logo_html}
          <div style="color:#ffffff;font-size:13px;opacity:0.9;margin-top:4px;">Recruitment System</div>
        </td></tr>
        <tr><td style="padding:28px 32px;color:#2d3748;font-size:15px;line-height:1.6;">{greeting_html}{body_html}{footer_html}</td></tr>
        <tr><td style="padding:16px 32px 24px;color:#888;font-size:12px;border-top:1px solid #eee;">
          {company} &mdash; Recruitment Team
        </td></tr>
      </table>
    </td></tr>
  </table>
</body></html>"""


def _email_history_file():
    return DATA_DIR / "email_history.json"


def save_email_history():
    serializable = []
    for entry in EMAIL_HISTORY[-2000:]:
        item = dict(entry)
        if isinstance(item.get("timestamp"), datetime):
            item["timestamp"] = item["timestamp"].isoformat()
        serializable.append(item)
    _save_json(_email_history_file(), serializable)


def log_email_history(recipient, subject, body, status, sender_id=None, application_id=None, email_type="general", error=None):
    import uuid
    sender = USERS.get(sender_id or "", {})
    entry = {
        "id": f"email-{str(uuid.uuid4())[:8]}",
        "sender_id": sender_id or "system",
        "sender_name": sender.get("name") or ("System" if sender_id == "system" else "Unknown"),
        "recipient": recipient,
        "subject": subject,
        "body_preview": (body or "")[:300],
        "application_id": application_id,
        "email_type": email_type,
        "status": status,
        "error": error,
        "timestamp": datetime.now(timezone.utc),
    }
    EMAIL_HISTORY.append(entry)
    save_email_history()


def send_email_message(
    recipient,
    subject,
    body,
    sender_id=None,
    application_id=None,
    email_type="general",
    candidate_name=None,
    footer_location=None,
):
    if not recipient:
        return False, "Candidate email is missing"

    mail = _mail_settings()
    if not mail["server"] or not mail["port"] or not mail["sender"]:
        log_email_history(recipient, subject, body, "failed", sender_id, application_id, email_type, "SMTP settings are incomplete")
        return False, "SMTP settings are incomplete"
    if not mail["username"] or not mail["password"]:
        log_email_history(recipient, subject, body, "failed", sender_id, application_id, email_type, "SMTP username or password is missing")
        return False, "SMTP username or password is missing — check .env (SMTP_USER, SMTP_PASS) or HR settings"

    plain_body = _format_email_content(body, candidate_name=candidate_name, footer_location=footer_location)
    logo_bytes, logo_subtype = _get_company_logo_for_email()
    use_cid_logo = bool(logo_bytes and logo_subtype)
    html_body = _build_email_html(
        body,
        candidate_name=candidate_name,
        footer_location=footer_location,
        use_cid_logo=use_cid_logo,
    )

    if use_cid_logo:
        message = MIMEMultipart("related")
        message["Subject"] = subject
        message["From"] = mail["sender"]
        message["To"] = recipient
        alt_part = MIMEMultipart("alternative")
        message.attach(alt_part)
        alt_part.attach(MIMEText(plain_body, "plain", "utf-8"))
        alt_part.attach(MIMEText(html_body, "html", "utf-8"))
        img = MIMEImage(logo_bytes, _subtype=logo_subtype)
        img.add_header("Content-ID", "<company_logo>")
        img.add_header("Content-Disposition", "inline", filename="company_logo.png")
        message.attach(img)
    else:
        message = EmailMessage()
        message["Subject"] = subject
        message["From"] = mail["sender"]
        message["To"] = recipient
        message.set_content(plain_body)
        message.add_alternative(html_body, subtype="html")

    try:
        if mail["port"] == 465:
            smtp_client = smtplib.SMTP_SSL(mail["server"], mail["port"], timeout=20)
        else:
            smtp_client = smtplib.SMTP(mail["server"], mail["port"], timeout=20)

        with smtp_client as smtp:
            if mail["port"] != 465 and mail["use_tls"]:
                smtp.starttls()
            if mail["username"]:
                smtp.login(mail["username"], mail["password"])
            smtp.send_message(message)

        log_email_history(recipient, subject, body, "success", sender_id, application_id, email_type)
        log_activity(sender_id or "system", "email_sent", f"Email sent to {recipient}: {subject}", "success")
        return True, None
    except Exception as exc:
        error_msg = str(exc)
        print(f"Warning: Could not send email: {error_msg}")
        log_email_history(recipient, subject, body, "failed", sender_id, application_id, email_type, error_msg)
        log_activity(sender_id or "system", "email_failed", f"Failed to send email to {recipient}: {error_msg}", "failed")
        return False, error_msg


def load_all_data():
    """Load all data from disk"""
    # Load users
    users_file = DATA_DIR / "users.json"
    if users_file.exists():
        try:
            with open(users_file, 'r') as f:
                USERS.update(json.load(f))
        except Exception as e:
            print(f"Warning: Could not load users: {e}")

    # Ensure default users exist
    if not USERS:
        USERS["user-001"] = {
            "id": "user-001",
            "email": "hr@company.com",
            "name": "HR Manager",
            "password_hash": generate_password_hash("password123"),
            "role": "hr"
        }
        USERS["user-002"] = {
            "id": "user-002",
            "email": "audit@company.com",
            "name": "Audit Manager",
            "password_hash": generate_password_hash("password123"),
            "role": "audit"
        }
        try:
            with open(users_file, 'w') as f:
                json.dump(USERS, f, indent=2, default=str)
        except Exception as e:
            print(f"Warning: Could not save default users: {e}")

    # Load jobs
    jobs_file = DATA_DIR / "jobs.json"
    if jobs_file.exists():
        try:
            with open(jobs_file, 'r') as f:
                JOBS.update(json.load(f))
        except Exception as e:
            print(f"Warning: Could not load jobs: {e}")
    for job in JOBS.values():
        if "positions_needed" not in job:
            job["positions_needed"] = 1

    # Load applications
    apps_file = DATA_DIR / "applications.json"
    if apps_file.exists():
        try:
            with open(apps_file, 'r') as f:
                APPLICATIONS.update(json.load(f))
        except Exception as e:
            print(f"Warning: Could not load applications: {e}")

    panels_file = DATA_DIR / "interviewer_panels.json"
    if panels_file.exists():
        try:
            with open(panels_file, "r") as f:
                INTERVIEWER_PANELS.update(json.load(f))
        except Exception as e:
            print(f"Warning: Could not load interviewer panels: {e}")

    # Load activity logs
    logs_file = DATA_DIR / "activity_logs.json"
    if logs_file.exists():
        try:
            with open(logs_file, 'r') as f:
                for entry in json.load(f):
                    if isinstance(entry.get('timestamp'), str):
                        try:
                            entry['timestamp'] = datetime.fromisoformat(entry['timestamp'])
                        except Exception:
                            entry['timestamp'] = datetime.now(timezone.utc)
                    ACTIVITY_LOGS.append(entry)
        except Exception as e:
            print(f"Warning: Could not load activity logs: {e}")

    # Backfill USERS schema
    for u in USERS.values():
        if "can_screen_interviews" not in u:
            u["can_screen_interviews"] = False
        if "can_score_interviews" not in u:
            u["can_score_interviews"] = False
        if "interviewer_token" not in u:
            u["interviewer_token"] = None
        if "interviewer_token_expires" not in u:
            u["interviewer_token_expires"] = None
        if "interviewer_token_used" not in u:
            u["interviewer_token_used"] = False
        if "interviewer_job_id" not in u:
            u["interviewer_job_id"] = None
        if "password_reset_required" not in u:
            u["password_reset_required"] = False
        if "first_login" not in u:
            u["first_login"] = True

    # Backfill APPLICATIONS schema
    for app_data in APPLICATIONS.values():
        if "interview_status" not in app_data:
            if app_data.get("status") == "interview":
                app_data["interview_status"] = "scheduled"
            else:
                app_data["interview_status"] = None
        if "interview_removed" not in app_data:
            app_data["interview_removed"] = False
        if "interview_removed_reason" not in app_data:
            app_data["interview_removed_reason"] = ""
        if "interview_scores" not in app_data:
            app_data["interview_scores"] = {
                "presentation": 0,
                "dressing": 0,
                "communication": 0,
                "confidence": 0,
                "technical_knowledge": 0,
                "problem_solving": 0,
                "attitude": 0,
            }
        else:
            scores = app_data["interview_scores"]
            for crit in ["presentation", "dressing", "communication", "confidence", "technical_knowledge", "problem_solving", "attitude"]:
                if crit not in scores:
                    scores[crit] = 0
        if "interview_total" not in app_data:
            app_data["interview_total"] = 0
        if "interview_scored_by" not in app_data:
            app_data["interview_scored_by"] = None
        if "interview_scored_at" not in app_data:
            app_data["interview_scored_at"] = None
        if "interview_score_submissions" not in app_data:
            app_data["interview_score_submissions"] = []
            if app_data.get("interview_scored_by") and app_data.get("interview_total"):
                scorer = USERS.get(app_data["interview_scored_by"], {})
                app_data["interview_score_submissions"].append({
                    "scorer_id": app_data["interview_scored_by"],
                    "scorer_name": scorer.get("name") or scorer.get("email") or "Unknown",
                    "scorer_role": scorer.get("role") or "",
                    "scores": dict(app_data.get("interview_scores") or {}),
                    "total": app_data.get("interview_total", 0),
                    "scored_at": app_data.get("interview_scored_at"),
                })
        if "selected" not in app_data:
            app_data["selected"] = False
        if "selected_at" not in app_data:
            app_data["selected_at"] = None
        if "selection_method" not in app_data:
            app_data["selection_method"] = ""

    settings_file = _settings_file()
    SETTINGS.update(DEFAULT_SETTINGS)
    if settings_file.exists():
        try:
            with open(settings_file, 'r') as f:
                SETTINGS.update(json.load(f))
        except Exception as e:
            print(f"Warning: Could not load settings: {e}")

    _apply_env_mail_settings()
    if not SETTINGS.get("mail_username") or not SETTINGS.get("mail_password"):
        print("Warning: SMTP credentials are not configured. Emails will not send until mail settings are saved.")

    if (SETTINGS.get("company_logo_data") or "").startswith("data:image") and not SETTINGS.get("company_logo_file"):
        _persist_company_logo_file(data_uri=SETTINGS["company_logo_data"])

    email_history_file = _email_history_file()
    if email_history_file.exists():
        try:
            with open(email_history_file, 'r') as f:
                for entry in json.load(f):
                    if isinstance(entry.get('timestamp'), str):
                        try:
                            entry['timestamp'] = datetime.fromisoformat(entry['timestamp'])
                        except Exception:
                            entry['timestamp'] = datetime.now(timezone.utc)
                    EMAIL_HISTORY.append(entry)
        except Exception as e:
            print(f"Warning: Could not load email history: {e}")

    panels_dirty = False
    users_dirty = False
    for user_id, interviewer in USERS.items():
        if interviewer.get("role") != "interviewer":
            continue
        job_id = interviewer.get("interviewer_job_id")
        if not job_id:
            continue
        if job_id not in INTERVIEWER_PANELS:
            INTERVIEWER_PANELS[job_id] = {
                "job_id": job_id,
                "token": None,
                "expires": None,
                "access_code": None,
                "active_application_id": None,
                "active_changed_at": None,
                "interviewer_ids": [],
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }
            panels_dirty = True
        panel = INTERVIEWER_PANELS[job_id]
        if "access_code" not in panel:
            panel["access_code"] = None
            panels_dirty = True
        if "active_application_id" not in panel:
            panel["active_application_id"] = None
            panels_dirty = True
        if "active_changed_at" not in panel:
            panel["active_changed_at"] = None
            panels_dirty = True
        if user_id not in panel.get("interviewer_ids", []):
            panel.setdefault("interviewer_ids", []).append(user_id)
            panels_dirty = True
        if interviewer.get("interviewer_token") and not panel.get("token"):
            panel["token"] = interviewer["interviewer_token"]
            panel["expires"] = interviewer.get("interviewer_token_expires")
            panels_dirty = True
        if interviewer.get("password_reset_required"):
            interviewer["password_reset_required"] = False
            users_dirty = True
    if panels_dirty or users_dirty:
        try:
            if panels_dirty:
                save_interviewer_panels()
            if users_dirty:
                save_users()
        except Exception as e:
            print(f"Warning: Could not migrate interviewer panels: {e}")


def log_activity(user_id, action, details="", status="success"):
    """Append an activity log entry and persist it."""
    entry = {
        "user_id": user_id,
        "action": action,
        "details": details,
        "status": status,
        "timestamp": datetime.now(timezone.utc),
    }
    ACTIVITY_LOGS.append(entry)
    try:
        serializable = []
        for e in ACTIVITY_LOGS[-1000:]:
            ec = dict(e)
            if isinstance(ec.get('timestamp'), datetime):
                ec['timestamp'] = ec['timestamp'].isoformat()
            serializable.append(ec)
        with open(DATA_DIR / "activity_logs.json", 'w') as f:
            json.dump(serializable, f, indent=2)
    except Exception as ex:
        print(f"Warning: Could not save activity logs: {ex}")


def get_current_user():
    """Get currently logged-in user"""
    if "user_id" in session:
        return USERS.get(session["user_id"])
    return None


def _portal_logo_src():
    logo_data = (SETTINGS.get("company_logo_data") or "").strip()
    if logo_data.startswith("data:image"):
        return logo_data
    return url_for("static", filename="gmb-logo.svg")


def _portal_welcome_image_src():
    image_data = (SETTINGS.get("portal_welcome_image_data") or "").strip()
    if image_data.startswith("data:image"):
        return image_data
    return ""


def _portal_welcome_video_src():
    video_name = (SETTINGS.get("portal_welcome_video") or "").strip()
    if not video_name:
        return ""
    video_path = BRANDING_DIR / video_name
    if video_path.exists():
        return url_for("portal_branding_file", filename=video_name)
    return ""


@app.context_processor
def inject_user():
    user = get_current_user()
    return {
        'current_user': user,
        'can_manage_users': bool(user and user.get('role') == 'hr'),
        'can_screen_interviews': bool(user and (user.get('role') == 'hr' or user.get('can_screen_interviews'))),
        'can_score_interviews': bool(user and (user.get('role') == 'hr' or user.get('can_score_interviews'))),
        'is_audit_manager': bool(user and user.get('role') == 'audit'),
        'portal_logo_src': _portal_logo_src(),
        'portal_welcome_image_src': _portal_welcome_image_src(),
        'portal_welcome_video_src': _portal_welcome_video_src(),
        'portal_welcome_video_name': (SETTINGS.get("portal_welcome_video") or "").strip(),
        'company_display_name': SETTINGS.get("company_name", "Grain Marketing Board"),
    }


@app.after_request
def add_security_headers(response):
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    return response


@app.route("/health")
def health():
    return jsonify({
        "status": "ok",
        "users": len(USERS),
        "jobs": len(JOBS),
        "applications": len(APPLICATIONS),
        "logs": len(ACTIVITY_LOGS),
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })


@app.route("/")
def index():
    user = get_current_user()
    if user:
        if user.get("role") in ("audit", "audit_member"):
            return redirect(url_for("audit_dashboard"))
        if user.get("role") in HR_ROLES:
            return redirect(url_for("hr_dashboard"))
        if user.get("role") == "interviewer":
            return redirect(url_for("interviewer_dashboard"))
    return redirect(url_for("login"))


@app.route("/portal")
def applicant_portal():
    """Public applicant job board — always shows open positions, even if HR is logged in."""
    return render_template("index.html", jobs=_open_jobs_for_portal())


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "").strip()

        user = None
        for u in USERS.values():
            if u.get("email", "").lower() == email:
                user = u
                break

        if user and check_password_hash(user.get("password_hash", ""), password):
            # Non-permanent session to clear on tab/window close
            session["user_id"] = user["id"]
            user["last_login"] = datetime.now(timezone.utc).isoformat()
            try:
                with open(DATA_DIR / "users.json", "w") as f:
                    json.dump(USERS, f, indent=2, default=str)
            except Exception as e:
                print(f"Warning: Could not save user login: {e}")
            log_activity(user["id"], "login", f"User {user.get('email')} logged in", "success")
            
            if user.get("role") in ("audit", "audit_member"):
                return redirect(url_for("audit_dashboard"))
            if user.get("role") in HR_ROLES:
                return redirect(url_for("hr_dashboard"))
            if user.get("role") == "interviewer":
                return redirect(url_for("interviewer_dashboard"))
                
            flash("Your account does not have dashboard access.", "warning")
            return redirect(url_for("login"))
        else:
            uid = user["id"] if user else "unknown"
            log_activity(uid, "login_failed", f"Failed login for {email}", "failed")
            flash("Invalid email or password.", "danger")

    return render_template("login.html")


@app.route("/logout")
def logout():
    user = get_current_user()
    if user:
        log_activity(user["id"], "logout", f"User {user.get('email')} logged out", "success")
    session.clear()
    flash("Logged out successfully.", "info")
    return redirect(url_for("login"))


@app.route("/hr")
def hr_dashboard():
    user = get_current_user()
    if not user:
        flash("Please log in.", "warning")
        return redirect(url_for("login"))
    if user.get("role") not in HR_ROLES:
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    _screen_due_jobs()

    dashboards = []
    for job_id, job in JOBS.items():
        applicants = [_enrich_application(a) for a in APPLICATIONS.values() if a.get("job_id") == job_id]
        enriched_job = _enrich_job(job)
        # Convert uploaded_at strings to datetime objects
        for a in applicants:
            if isinstance(a.get("uploaded_at"), str):
                try:
                    a["uploaded_at"] = datetime.fromisoformat(a["uploaded_at"])
                except Exception:
                    a["uploaded_at"] = datetime.now(timezone.utc)
        dashboards.append({
            "job": enriched_job,
            "applicants": applicants,
            "applicant_count": len(applicants),
            "all_count": len(applicants),
            "selected_count": len([a for a in applicants if a.get("status") == "shortlisted"]),
            "deadline_state": enriched_job.get("deadline_state", "open"),
        })

    all_applicants_data = [_enrich_application(app_data) for app_data in APPLICATIONS.values()]
    # Convert uploaded_at strings to datetime objects and sort newest first
    for a in all_applicants_data:
        if isinstance(a.get("uploaded_at"), str):
            try:
                a["uploaded_at"] = datetime.fromisoformat(a["uploaded_at"])
            except Exception:
                a["uploaded_at"] = datetime.now(timezone.utc)
    all_applicants_data.sort(key=lambda item: item.get("uploaded_at") or datetime.min.replace(tzinfo=timezone.utc), reverse=True)

    email_history = sorted(
        [dict(entry) for entry in EMAIL_HISTORY],
        key=lambda item: item.get("timestamp") or datetime.min.replace(tzinfo=timezone.utc),
        reverse=True,
    )
    for entry in email_history:
        if isinstance(entry.get("timestamp"), datetime):
            entry["timestamp"] = entry["timestamp"]
    
    # Create jobs list with application counts for job filter
    jobs_list = []
    for job_id, job in JOBS.items():
        app_count = len([a for a in APPLICATIONS.values() if a.get("job_id") == job_id])
        jobs_list.append({
            "id": job_id,
            "title": job.get("title", "Unknown Job"),
            "applications_count": app_count,
            "positions_needed": job.get("positions_needed", 1),
        })

    interview_candidates = [
        _enrich_application(a)
        for a in APPLICATIONS.values()
        if a.get("status") in ("shortlisted", "interview", "pending", "selected")
    ]
    selected_by_job = {}
    for app_data in APPLICATIONS.values():
        if app_data.get("selected"):
            enriched = _enrich_application(app_data)
            job_title = enriched.get("job_title", "Unknown Job")
            selected_by_job.setdefault(job_title, []).append(enriched)
    
    interviews_data = _scheduled_interviews()
    interview_marks_data = _interview_marks_board()

    return render_template(
        "hr_dashboard.html",
        dashboards=dashboards,
        dashboards_json=_json_safe(dashboards),
        applicants=all_applicants_data,
        applicants_json=_json_safe(all_applicants_data),
        all_applicants=all_applicants_data,
        recent_applicants=all_applicants_data[:5],
        interviews=interviews_data,
        interviews_json=_json_safe(interviews_data),
        interview_candidates=interview_candidates,
        interview_candidates_json=_json_safe(interview_candidates),
        interview_marks_board=interview_marks_data,
        interview_marks_board_json=_json_safe(interview_marks_data),
        selected_by_job=selected_by_job,
        interviewer_panels=_interviewer_panels_for_dashboard(),
        settings=SETTINGS,
        jobs=jobs_list,
        email_history=email_history,
        can_manage_users=user.get("role") == "hr",
        users=list(USERS.values()),
        hr_notifications=_build_hr_notifications(),
    )


@app.route("/audit")
def audit_dashboard():
    user = get_current_user()
    if not user:
        flash("Please log in.", "warning")
        return redirect(url_for("login"))
    if user.get("role") not in ("audit", "audit_member"):
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    PER_PAGE = 25
    page = max(1, request.args.get("page", 1, type=int))
    action_filter = request.args.get("action", "").strip().lower()

    def _normalize_log(entry):
        item = dict(entry or {})
        ts = item.get("timestamp")
        if isinstance(ts, str):
            try:
                ts = datetime.fromisoformat(ts)
            except Exception:
                ts = datetime.now(timezone.utc)
        if not isinstance(ts, datetime):
            ts = datetime.now(timezone.utc)

        user_ref = USERS.get(item.get("user_id"), {})
        item["timestamp"] = ts
        item["created_at"] = ts
        item["user"] = {
            "full_name": user_ref.get("name") or item.get("user_id") or "Unknown User",
            "email": user_ref.get("email") or "unknown@local",
        }
        return item

    # Sort newest first
    all_logs = sorted([_normalize_log(log) for log in ACTIVITY_LOGS], key=lambda x: x.get("timestamp"), reverse=True)

    filtered = [l for l in all_logs if action_filter in l.get("action", "").lower()] if action_filter else all_logs

    total_logs = len(all_logs)
    successful_count = sum(1 for l in all_logs if l.get("status") == "success")
    failed_count = sum(1 for l in all_logs if l.get("status") != "success")
    unique_users = len(set(l.get("user_id") for l in all_logs if l.get("user_id")))

    login_logs = [l for l in all_logs if l.get("action") in ("login", "logout")][:20]

    total_pages = max(1, (len(filtered) + PER_PAGE - 1) // PER_PAGE)
    page = min(page, total_pages)
    logs = filtered[(page - 1) * PER_PAGE: page * PER_PAGE]

    audit_members = [u for u in USERS.values() if u.get("role") == "audit_member"]

    return render_template(
        "audit_dashboard.html",
        logs=logs,
        login_logs=login_logs,
        login_activities=login_logs,
        total_logs=total_logs,
        successful_count=successful_count,
        failed_count=failed_count,
        unique_users=unique_users,
        page=page,
        current_page=page,
        total_pages=total_pages,
        action_filter=action_filter,
        audit_members=audit_members,
    )


@app.route("/audit/users/add", methods=["POST"])
def audit_add_member():
    user = get_current_user()
    if not user or user.get("role") != "audit":
        flash("Unauthorized", "danger")
        return redirect(url_for("login"))

    name = request.form.get("name", "").strip()
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "").strip()

    if not name or not email or not password:
        flash("All fields are required.", "danger")
        return redirect(url_for("audit_dashboard"))

    # Check if email exists
    for u in USERS.values():
        if u.get("email", "").lower() == email:
            flash("Email already exists.", "danger")
            return redirect(url_for("audit_dashboard"))

    import uuid
    new_user_id = f"user-{str(uuid.uuid4())[:8]}"
    new_user = {
        "id": new_user_id,
        "email": email,
        "name": name,
        "password_hash": generate_password_hash(password),
        "role": "audit_member",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "can_screen_interviews": False,
        "can_score_interviews": False,
        "interviewer_token": None,
        "interviewer_token_expires": None,
        "interviewer_token_used": False,
        "interviewer_job_id": None,
        "password_reset_required": False,
        "first_login": True
    }

    USERS[new_user_id] = new_user
    save_users()
    log_activity(user["id"], "add_audit_member", f"Added audit member: {email}", "success")
    flash(f"Audit member {name} added successfully.", "success")
    return redirect(url_for("audit_dashboard"))


# ─── Helper: dict wrapper for attribute access in templates ───────────────
class _DictObj:
    """Wrap a plain dict so Jinja2 can access fields as attributes,
    with special handling for file paths and datetimes."""
    def __init__(self, d):
        object.__setattr__(self, '_d', d or {})

    def __getattr__(self, name):
        d = object.__getattribute__(self, '_d')
        val = d.get(name)
        if name == 'resume_path' and val:
            return Path(val)
        if name == 'certification_paths':
            return [Path(p) for p in (val or [])]
        if name == 'uploaded_at' and isinstance(val, str):
            try:
                return datetime.fromisoformat(val)
            except Exception:
                return datetime.now(timezone.utc)
        if name in ('due_date', 'interview_at') and isinstance(val, str):
            parsed = _parse_iso_datetime(val)
            if parsed:
                return parsed
        if name == 'requirements':
            return _normalize_requirements(val)
        return val

    def __getitem__(self, key):
        return object.__getattribute__(self, '_d').get(key)

    def __bool__(self):
        return bool(object.__getattribute__(self, '_d'))

    def get(self, key, default=None):
        """Dict-like get method"""
        return object.__getattribute__(self, '_d').get(key, default)


class _AppObj(_DictObj):
    @property
    def job_ref(self):
        job = JOBS.get(object.__getattribute__(self, '_d').get('job_id', ''))
        return _DictObj(job) if job else None


def _wrap_app(d):
    return _AppObj(d)


def _wrap_job(d):
    return _DictObj(d)


def _json_safe(obj):
    if isinstance(obj, datetime):
        return obj.isoformat()
    if isinstance(obj, dict):
        return {key: _json_safe(value) for key, value in obj.items()}
    if isinstance(obj, list):
        return [_json_safe(value) for value in obj]
    return obj


def _enrich_job(job):
    item = dict(job or {})
    item["requirements"] = _normalize_requirements(item.get("requirements"))
    due_date = _parse_iso_datetime(item.get("due_date"))
    item["due_date"] = due_date.isoformat() if due_date else item.get("due_date")
    item["application_link"] = item.get("application_link") or url_for("apply", job_id=item.get("id", ""))
    item["deadline_state"] = "closed" if due_date and due_date < datetime.now(timezone.utc) else "open"
    try:
        item["positions_needed"] = max(1, int(item.get("positions_needed") or 1))
    except (TypeError, ValueError):
        item["positions_needed"] = 1
    return item


def _has_application_for_job(job_id, email):
    email_key = (email or "").strip().lower()
    if not email_key:
        return False
    return any(
        app.get("job_id") == job_id and (app.get("email") or "").strip().lower() == email_key
        for app in APPLICATIONS.values()
    )


def _resolve_upload_path(path_str):
    if not path_str:
        return None
    path = Path(path_str)
    if path.exists():
        return path
    upload_dir = Path(__file__).parent / "uploads"
    candidate = upload_dir / path.name
    if candidate.exists():
        return candidate
    return None


def _enrich_application(app_data):
    item = dict(app_data or {})
    job = JOBS.get(item.get("job_id", ""), {})
    item["display_name"] = _candidate_full_name(item)
    item["job_title"] = job.get("title") or "General Application"
    job_due_date = _parse_iso_datetime(job.get("due_date"))
    item["job_due_date"] = job_due_date.isoformat() if job_due_date else ""
    item["job_deadline_state"] = "closed" if job_due_date and job_due_date <= datetime.now(timezone.utc) else "open"
    interview_date = (item.get("interview_date") or "").strip()
    interview_time = (item.get("interview_time") or "").strip()
    if interview_date and interview_time:
        item["interview_at"] = f"{interview_date}T{interview_time}"
    elif interview_date:
        item["interview_at"] = interview_date

    resume_path = _resolve_upload_path(item.get("resume_path"))
    item["resume_basename"] = resume_path.name if resume_path else ""
    app_id = item.get("id")
    if app_id and resume_path:
        item["resume_view_url"] = url_for("hr_application_resume", application_id=app_id)
        item["resume_download_url"] = url_for("hr_application_resume", application_id=app_id, download=1)

    cert_files = []
    for index, cert_ref in enumerate(item.get("certification_paths") or []):
        cert_path = _resolve_upload_path(cert_ref)
        if not cert_path or not app_id:
            continue
        cert_files.append({
            "index": index,
            "name": cert_path.name,
            "view_url": url_for("hr_application_certificate", application_id=app_id, cert_index=index),
            "download_url": url_for("hr_application_certificate", application_id=app_id, cert_index=index, download=1),
        })
    item["certificate_files"] = cert_files
    item["score_submissions"] = _enrich_interview_score_submissions(item)
    item["interview_average"] = _interview_score_average(item)
    required_scorers = _required_scorer_ids(item.get("job_id"))
    submitted_scorers = _submitted_scorer_ids(item)
    item["required_scorer_count"] = len(required_scorers) if required_scorers else 1
    item["submitted_scorer_count"] = len(submitted_scorers)
    item["all_scorers_submitted"] = _all_required_scorers_submitted(item)
    item["pending_scorer_names"] = _pending_scorer_names(item)
    item["is_active_interview"] = _is_active_interview_candidate(item)
    return item


INTERVIEW_SCORE_CRITERIA = [
    "presentation", "dressing", "communication", "confidence",
    "technical_knowledge", "problem_solving", "attitude",
]


def _enrich_interview_score_submissions(app_data):
    submissions = []
    for entry in app_data.get("interview_score_submissions") or []:
        item = dict(entry or {})
        scorer = USERS.get(item.get("scorer_id"), {})
        item["scorer_name"] = item.get("scorer_name") or scorer.get("name") or scorer.get("email") or "Unknown"
        item["scorer_role"] = item.get("scorer_role") or scorer.get("role") or ""
        if item.get("scored_at") and isinstance(item["scored_at"], str):
            try:
                item["scored_at_display"] = datetime.fromisoformat(item["scored_at"]).strftime("%Y-%m-%d %H:%M")
            except Exception:
                item["scored_at_display"] = item["scored_at"][:16]
        else:
            item["scored_at_display"] = ""
        submissions.append(item)
    submissions.sort(key=lambda s: s.get("scored_at") or "", reverse=True)
    return submissions


def _interview_score_average(app_data):
    submissions = app_data.get("interview_score_submissions") or []
    if not submissions:
        return app_data.get("interview_total") or 0
    totals = [float(s.get("total") or 0) for s in submissions]
    if not totals:
        return 0
    return round(sum(totals) / len(totals), 2)


def _upsert_interview_score_submission(app_data, user, scores, percentage):
    submissions = app_data.setdefault("interview_score_submissions", [])
    entry = {
        "scorer_id": user["id"],
        "scorer_name": user.get("name") or user.get("email") or "Unknown",
        "scorer_role": user.get("role") or "",
        "scores": scores,
        "total": round(percentage, 2),
        "scored_at": datetime.now(timezone.utc).isoformat(),
    }
    replaced = False
    for index, existing in enumerate(submissions):
        if existing.get("scorer_id") == user["id"]:
            submissions[index] = entry
            replaced = True
            break
    if not replaced:
        submissions.append(entry)

    app_data["interview_total"] = _interview_score_average(app_data)
    app_data["interview_scores"] = scores
    app_data["interview_scored_by"] = user["id"]
    app_data["interview_scored_at"] = entry["scored_at"]
    return entry


def _required_scorer_ids(job_id):
    """External interviewers assigned to score candidates for a job."""
    panel = INTERVIEWER_PANELS.get(job_id, {})
    return list(panel.get("interviewer_ids") or [])


def _submitted_scorer_ids(app_data):
    return {
        entry.get("scorer_id")
        for entry in (app_data.get("interview_score_submissions") or [])
        if entry.get("scorer_id")
    }


def _all_required_scorers_submitted(app_data):
    required = _required_scorer_ids(app_data.get("job_id"))
    submitted = _submitted_scorer_ids(app_data)
    if not required:
        return len(submitted) > 0
    return all(scorer_id in submitted for scorer_id in required)


def _sync_interview_status_after_score(app_data):
    """Keep interviews open until every required evaluator has scored and HR screens."""
    if app_data.get("interview_status") in ("completed", "cancelled"):
        return
    if _all_required_scorers_submitted(app_data):
        app_data["interview_status"] = "awaiting_screening"
    elif app_data.get("interview_score_submissions"):
        if app_data.get("interview_status") != "in_progress":
            app_data["interview_status"] = "in_progress"
    elif app_data.get("interview_status") not in ("scheduled", "in_progress", "awaiting_screening"):
        app_data["interview_status"] = "in_progress"


def _get_panel_active_application_id(job_id):
    panel = INTERVIEWER_PANELS.get(job_id, {})
    return panel.get("active_application_id")


def _is_active_interview_candidate(app_data):
    if not app_data:
        return False
    job_id = app_data.get("job_id")
    return (
        _get_panel_active_application_id(job_id) == app_data.get("id")
        and app_data.get("interview_status") == "in_progress"
    )


def _pending_scorer_names(app_data):
    required = _required_scorer_ids(app_data.get("job_id"))
    submitted = _submitted_scorer_ids(app_data)
    pending = []
    for user_id in required:
        if user_id in submitted:
            continue
        user = USERS.get(user_id, {})
        pending.append(user.get("name") or user.get("email") or "Unknown interviewer")
    return pending


def _interview_queue_for_job(job_id):
    apps = []
    for app_data in APPLICATIONS.values():
        if app_data.get("job_id") != job_id or app_data.get("interview_removed"):
            continue
        if app_data.get("status") not in ("shortlisted", "interview", "selected"):
            continue
        if not (app_data.get("interview_date") or app_data.get("interview_status")):
            continue
        apps.append(app_data)
    apps.sort(
        key=lambda item: (
            item.get("interview_date") or "9999",
            item.get("interview_time") or "",
            item.get("name") or "",
        )
    )
    return apps


def _set_active_interview(job_id, app_id):
    panel = _get_interviewer_panel(job_id, create=True)
    for other in _interview_queue_for_job(job_id):
        if other.get("id") == app_id:
            continue
        if other.get("interview_status") == "in_progress":
            _sync_interview_status_after_score(other)
            if other.get("interview_status") == "in_progress":
                if _all_required_scorers_submitted(other):
                    other["interview_status"] = "awaiting_screening"
                else:
                    other["interview_status"] = "awaiting_screening"
    panel["active_application_id"] = app_id
    panel["active_changed_at"] = datetime.now(timezone.utc).isoformat()
    panel["updated_at"] = datetime.now(timezone.utc).isoformat()


def _interview_marks_board():
    board = []
    for app_data in APPLICATIONS.values():
        if app_data.get("interview_removed"):
            continue
        status = app_data.get("interview_status")
        has_schedule = bool((app_data.get("interview_date") or "").strip())
        has_marks = bool(app_data.get("interview_score_submissions"))
        if not has_schedule and status not in ("scheduled", "in_progress", "awaiting_screening", "completed") and not has_marks:
            continue
        board.append(_enrich_application(app_data))
    board.sort(
        key=lambda item: (
            item.get("interview_date") or "9999",
            item.get("interview_time") or "",
            item.get("display_name") or "",
        ),
    )
    return board


def _open_jobs_for_portal():
    jobs = []
    for job_id, job_data in JOBS.items():
        enriched = _enrich_job(job_data)
        if enriched.get("deadline_state") != "open":
            continue
        wrapped = _wrap_job(enriched)
        app_count = len([a for a in APPLICATIONS.values() if a.get("job_id") == job_id])
        wrapped._d["applications_count"] = app_count
        jobs.append(wrapped)
    return jobs


def _build_hr_notifications():
    notifications = []
    now = datetime.now(timezone.utc)

    pending_apps = sorted(
        [a for a in APPLICATIONS.values() if str(a.get("status") or "").lower() == "pending"],
        key=lambda item: item.get("uploaded_at") or "",
        reverse=True,
    )
    for app_data in pending_apps[:8]:
        enriched = _enrich_application(app_data)
        notifications.append({
            "id": f"app-{app_data.get('id', '')}",
            "type": "application",
            "title": "New Application",
            "message": f"{_candidate_full_name(app_data)} applied for {enriched.get('job_title', 'a role')}",
            "action": "applicants",
            "target_type": "application",
            "target_id": app_data.get("id", ""),
        })

    for interview in _scheduled_interviews():
        interview_date = (interview.get("interview_date") or "").strip()
        if not interview_date:
            continue
        try:
            interview_dt = datetime.fromisoformat(f"{interview_date}T{(interview.get('interview_time') or '09:00')[:5]}")
            if interview_dt.tzinfo is None:
                interview_dt = interview_dt.replace(tzinfo=timezone.utc)
        except Exception:
            continue
        days_left = (interview_dt.date() - now.date()).days
        if 0 <= days_left <= 7:
            notifications.append({
                "id": f"interview-{interview.get('id', '')}-{interview_date}",
                "type": "interview",
                "title": "Upcoming Interview",
                "message": f"{_candidate_full_name(interview)} — {interview_date} at {interview.get('interview_time', '')}",
                "action": "interviews",
                "target_type": "application",
                "target_id": interview.get("id", ""),
            })

    for job_id, job in JOBS.items():
        due_date = _parse_iso_datetime(job.get("due_date"))
        if not due_date or due_date <= now:
            continue
        days_left = (due_date.date() - now.date()).days
        if days_left <= 7:
            notifications.append({
                "id": f"deadline-{job_id}-{due_date.strftime('%Y%m%d')}",
                "type": "deadline",
                "title": "Application Deadline Approaching",
                "message": f"{job.get('title', 'Job')} closes on {due_date.strftime('%Y-%m-%d')} ({days_left} day(s) left)",
                "action": "jobs",
                "target_type": "job",
                "target_id": job_id,
            })

    return notifications


def _scheduled_interviews():
    interviews = []
    for app_data in APPLICATIONS.values():
        if app_data.get("interview_removed"):
            continue
        if app_data.get("interview_date"):
            enriched = _enrich_application(app_data)
            # Ensure interview_at is a datetime object
            if isinstance(enriched.get("interview_at"), str):
                try:
                    enriched["interview_at"] = datetime.fromisoformat(enriched["interview_at"])
                except Exception:
                    pass
            interviews.append(enriched)
    interviews.sort(key=lambda item: (item.get("interview_date", ""), item.get("interview_time", "")))
    return interviews


def _require_hr():
    user = get_current_user()
    if not user:
        flash("Please log in.", "warning")
        return None, redirect(url_for("login"))
    if user.get("role") not in HR_ROLES:
        flash("Access denied.", "danger")
        return None, redirect(url_for("login"))
    return user, None


def _require_hr_api():
    user = get_current_user()
    if not user:
        return None, (jsonify({"ok": False, "error": "Please log in."}), 401)
    if user.get("role") not in HR_ROLES:
        return None, (jsonify({"ok": False, "error": "Access denied."}), 403)
    return user, None


def _require_hr_admin():
    user = get_current_user()
    if not user:
        return None, (jsonify({"ok": False, "error": "Unauthorized"}), 401)
    if user.get("role") != "hr":
        return None, (jsonify({"ok": False, "error": "Only HR administrators can perform this action"}), 403)
    return user, None


# ─── HR sub-pages ──────────────────────────────────────────────────────────
@app.route("/hr/jobs", methods=["GET"])
def hr_jobs():
    user, err = _require_hr()
    if err:
        return err
    _screen_due_jobs()
    jobs = [_wrap_job(_enrich_job(j)) for j in JOBS.values()]
    return render_template("hr_jobs.html", jobs=jobs)


@app.route("/hr/jobs/post", methods=["POST"])
def hr_post_job():
    user, err = _require_hr()
    if err:
        return err
    import uuid
    job_id = str(uuid.uuid4())
    try:
        positions_needed = max(1, int(request.form.get("positions_needed", 1)))
    except (TypeError, ValueError):
        positions_needed = 1
    JOBS[job_id] = {
        "id": job_id,
        "title": request.form.get("title", "").strip(),
        "department": request.form.get("department", "").strip(),
        "location": request.form.get("location", "").strip(),
        "description": request.form.get("description", "").strip(),
        "requirements": _normalize_requirements(request.form.get("requirements", "")),
        "screening_criteria": request.form.get("screening_criteria", "").strip(),
        "due_date": request.form.get("due_date", "").strip(),
        "positions_needed": positions_needed,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "application_link": url_for("apply", job_id=job_id),
    }
    try:
        save_jobs()
    except Exception as e:
        print(f"Warning: Could not save job: {e}")
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True, "job": _enrich_job(JOBS[job_id])})
    flash("Job posted successfully.", "success")
    return redirect(url_for("hr_jobs"))


@app.route("/hr/screening")
def hr_screening():
    user, err = _require_hr()
    if err:
        return err
    screening_summary = _screen_due_jobs()
    all_apps = [_wrap_app(_enrich_application(a)) for a in APPLICATIONS.values()]
    pending = [a for a in all_apps if not a.status or a.status == "pending"]
    shortlisted = [a for a in all_apps if a.status == "shortlisted"]
    rejected = [a for a in all_apps if a.status == "rejected"]
    return render_template(
        "hr_screening.html",
        applicants=all_apps,
        pending=pending,
        shortlisted=shortlisted,
        rejected=rejected,
        screening_summary=screening_summary,
    )


@app.route("/hr/screening/run-now", methods=["POST"])
def hr_screening_run_now():
    user, err = _require_hr()
    if err:
        return err
    
    data = request.get_json() or {}
    job_id = data.get("job_id")
    
    # If a specific job is selected, screen only that job; otherwise screen all jobs
    job_ids = [job_id] if job_id else list(JOBS.keys())
    
    summary = _screen_pending_applications(job_ids, trigger="manual")
    job_titles = [JOBS.get(jid, {}).get("title", "Unknown Job") for jid in job_ids]
    action_msg = f"Processed {summary['processed']} candidate(s) for {', '.join(job_titles)}"
    log_activity(user["id"], "screen_due_candidates", action_msg, "success")
    return jsonify({"ok": True, **summary})


@app.route("/hr/interviews")
def hr_interviews():
    user, err = _require_hr()
    if err:
        return err
    return redirect(url_for("hr_dashboard") + "#interviews")


@app.route("/hr/ratings")
def hr_ratings():
    user, err = _require_hr()
    if err:
        return err
    applicants = [_wrap_app(a) for a in APPLICATIONS.values()]
    return render_template("hr_ratings.html", applicants=applicants)


@app.route("/hr/settings", methods=["GET"])
def hr_settings():
    user, err = _require_hr()
    if err:
        return err
    return render_template("hr_settings.html", settings=SETTINGS)


@app.route("/portal/branding/<filename>")
def portal_branding_file(filename):
    from werkzeug.utils import secure_filename
    safe_name = secure_filename(filename)
    target = BRANDING_DIR / safe_name
    if not target.exists():
        return "File not found", 404
    from flask import send_from_directory
    return send_from_directory(str(BRANDING_DIR), safe_name)


@app.route("/hr/settings/logo", methods=["POST"])
def hr_settings_logo():
    user, err = _require_hr()
    if err:
        return err
    payload = request.get_json(silent=True) or {}
    logo_data = (payload.get("logo") or "").strip()
    if logo_data and not logo_data.startswith("data:image"):
        return jsonify({"ok": False, "error": "Invalid logo format"}), 400
    SETTINGS["company_logo_data"] = logo_data
    if logo_data:
        _persist_company_logo_file(data_uri=logo_data)
    save_settings()
    return jsonify({"ok": True})


@app.route("/hr/settings/portal-branding", methods=["POST"])
def hr_settings_portal_branding():
    user, err = _require_hr()
    if err:
        return err
    from werkzeug.utils import secure_filename

    if request.is_json:
        payload = request.get_json(silent=True) or {}
        image_data = (payload.get("welcome_image") or "").strip()
        if image_data and not image_data.startswith("data:image"):
            return jsonify({"ok": False, "error": "Invalid image format"}), 400
        if image_data:
            SETTINGS["portal_welcome_image_data"] = image_data
        save_settings()
        return jsonify({"ok": True})

    company_logo = request.files.get("company_logo")
    if company_logo and company_logo.filename:
        logo_bytes = company_logo.read()
        mime = company_logo.mimetype or "image/png"
        SETTINGS["company_logo_data"] = (
            f"data:{mime};base64,{base64.b64encode(logo_bytes).decode('ascii')}"
        )
        _persist_company_logo_file(raw_bytes=logo_bytes, mime=mime)

    welcome_image = request.files.get("welcome_image")
    if welcome_image and welcome_image.filename:
        image_data = welcome_image.read()
        mime = welcome_image.mimetype or "image/png"
        SETTINGS["portal_welcome_image_data"] = (
            f"data:{mime};base64,{base64.b64encode(image_data).decode('ascii')}"
        )

    welcome_video = request.files.get("welcome_video")
    if welcome_video and welcome_video.filename:
        video_name = secure_filename(welcome_video.filename)
        save_path = BRANDING_DIR / video_name
        welcome_video.save(str(save_path))
        SETTINGS["portal_welcome_video"] = video_name

    save_settings()
    return jsonify({
        "ok": True,
        "portal_logo_src": _portal_logo_src(),
        "portal_welcome_image_src": _portal_welcome_image_src(),
        "portal_welcome_video_src": _portal_welcome_video_src(),
        "portal_welcome_video_name": (SETTINGS.get("portal_welcome_video") or "").strip(),
    })


@app.route("/hr/settings/test-email", methods=["POST"])
def hr_settings_test_email():
    user, err = _require_hr()
    if err:
        return err
    payload = request.get_json(silent=True) or {}
    recipient = (payload.get("email") or user.get("email") or "").strip()
    if not recipient:
        return jsonify({"ok": False, "error": "No recipient email provided"}), 400
    sent, error = send_email_message(
        recipient,
        "GMB Recruitment — Test Email",
        "This is a test email from the recruitment system. If you received this, SMTP is configured correctly.",
        sender_id=user["id"],
        email_type="test",
        candidate_name=user.get("name") or "HR User",
    )
    if sent:
        return jsonify({"ok": True, "message": f"Test email sent to {recipient}"})
    return jsonify({"ok": False, "error": error or "Failed to send test email"}), 500


@app.route("/hr/settings/update", methods=["POST"])
def hr_settings_update():
    user, err = _require_hr()
    if err:
        return err
    SETTINGS.update({
        "company_name": request.form.get("company_name", SETTINGS.get("company_name", DEFAULT_SETTINGS["company_name"])).strip(),
        "support_email": request.form.get("support_email", SETTINGS.get("support_email", DEFAULT_SETTINGS["support_email"])).strip(),
        "default_view": request.form.get("default_view", SETTINGS.get("default_view", DEFAULT_SETTINGS["default_view"])).strip(),
        "items_per_page": int(request.form.get("items_per_page", SETTINGS.get("items_per_page", DEFAULT_SETTINGS["items_per_page"]))),
        "notification_email_enabled": "notification_email_enabled" in request.form,
        "application_updates_enabled": "application_updates_enabled" in request.form,
        "interview_reminders_enabled": "interview_reminders_enabled" in request.form,
        "weekly_reports_enabled": "weekly_reports_enabled" in request.form,
        "mail_server": request.form.get("mail_server", SETTINGS.get("mail_server", DEFAULT_SETTINGS["mail_server"])).strip(),
        "mail_port": int(request.form.get("mail_port", SETTINGS.get("mail_port", DEFAULT_SETTINGS["mail_port"]))),
        "mail_use_tls": "mail_use_tls" in request.form,
        "mail_username": request.form.get("mail_username", SETTINGS.get("mail_username", DEFAULT_SETTINGS["mail_username"])).strip(),
        "mail_default_sender": request.form.get("mail_default_sender", SETTINGS.get("mail_default_sender", DEFAULT_SETTINGS["mail_default_sender"])).strip(),
    })
    # Only update password if a new one was provided (don't clear existing password)
    new_password = request.form.get("mail_password", "").strip()
    if new_password:
        SETTINGS["mail_password"] = new_password
    save_settings()
    flash("Settings saved.", "success")
    return redirect(url_for("hr_settings"))


@app.route("/hr/users/add", methods=["POST"])
def hr_add_user():
    user, err = _require_hr_admin()
    if err:
        return err
    
    name = request.form.get("name", "").strip()
    email = request.form.get("email", "").strip().lower()
    role = request.form.get("role", "hr_staff").strip()
    password = request.form.get("password", "").strip()
    
    if not name or not email or not password:
        return jsonify({"ok": False, "error": "Missing required fields"}), 400
    
    # Check if user already exists
    for u in USERS.values():
        if u.get("email", "").lower() == email:
            return jsonify({"ok": False, "error": "User with this email already exists"}), 400
    
    # Validate role - added users get limited HR access by default
    valid_roles = ["hr_staff", "interviewer"]
    if role not in valid_roles:
        role = "hr_staff"
    
    # Create new user
    import uuid
    new_user_id = f"user-{str(uuid.uuid4())[:8]}"
    new_user = {
        "id": new_user_id,
        "email": email,
        "name": name,
        "password_hash": generate_password_hash(password),
        "role": role,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "can_screen_interviews": False,
        "can_score_interviews": False,
        "interviewer_token": None,
        "interviewer_token_expires": None,
        "interviewer_token_used": False,
        "interviewer_job_id": None,
        "password_reset_required": False,
        "first_login": True
    }
    
    USERS[new_user_id] = new_user

    # Save to disk
    try:
        with open(DATA_DIR / "users.json", "w") as f:
            json.dump(USERS, f, indent=2, default=str)
    except Exception as e:
        print(f"Warning: Could not save users: {e}")
        return jsonify({"ok": False, "error": "Could not save user"}), 500
    log_activity(user["id"], "add_user", f"Added new user: {email} ({role})", "success")
    return jsonify({"ok": True, "message": f"User {name} added successfully with role {role}"})


@app.route("/hr/users/<user_id>/permissions", methods=["POST"])
def hr_toggle_user_permissions(user_id):
    user, err = _require_hr_admin()
    if err:
        return err
        
    u = USERS.get(user_id)
    if not u:
        return jsonify({"ok": False, "error": "User not found"}), 404
        
    payload = request.get_json(silent=True) or request.form
    u["can_screen_interviews"] = _to_bool(payload.get("can_screen_interviews", False))
    u["can_score_interviews"] = _to_bool(payload.get("can_score_interviews", False))
    
    save_users()
    log_activity(user["id"], "update_permissions", f"Updated permissions for user {u.get('email')}", "success")
    return jsonify({"ok": True, "message": "Permissions updated successfully"})


def _get_interviewer_panel(job_id, create=False):
    panel = INTERVIEWER_PANELS.get(job_id)
    if panel or not create:
        return panel
    panel = {
        "job_id": job_id,
        "token": None,
        "expires": None,
        "access_code": None,
        "active_application_id": None,
        "active_changed_at": None,
        "interviewer_ids": [],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    INTERVIEWER_PANELS[job_id] = panel
    return panel


def _get_panel_by_token(token):
    for panel in INTERVIEWER_PANELS.values():
        if panel.get("token") == token:
            return panel
    return None


def _panel_link_expired(panel):
    expires_str = panel.get("expires")
    if not expires_str:
        return True
    try:
        expires = datetime.fromisoformat(expires_str)
    except Exception:
        return True
    return datetime.now(timezone.utc) > expires


def _normalize_person_name(value):
    return " ".join((value or "").strip().lower().split())


def _clean_name_input(value):
    honorifics = {"dr", "mr", "mrs", "ms", "prof", "sir", "miss"}
    parts = _normalize_person_name(value).split()
    while parts and parts[0].rstrip(".") in honorifics:
        parts.pop(0)
    return " ".join(parts)


def _normalize_access_code(value):
    return re.sub(r"\D", "", (value or "").strip())


def _sync_panel_membership(panel):
    """Ensure every interviewer assigned to this job is on the panel list."""
    job_id = panel.get("job_id")
    if not job_id:
        return
    panel.setdefault("interviewer_ids", [])
    dirty = False
    for user in USERS.values():
        if user.get("role") != "interviewer":
            continue
        if user.get("interviewer_job_id") != job_id:
            continue
        if user["id"] not in panel["interviewer_ids"]:
            panel["interviewer_ids"].append(user["id"])
            dirty = True
    if dirty:
        panel["updated_at"] = datetime.now(timezone.utc).isoformat()
        save_interviewer_panels()


def _panel_interviewer_users(panel):
    members = []
    for user_id in panel.get("interviewer_ids") or []:
        user = USERS.get(user_id)
        if not user or user.get("role") != "interviewer":
            continue
        members.append(user)
    return members


def _find_panel_interviewer(panel, name_or_email):
    """Match interviewer by email, name, first name, or email username on this panel."""
    raw = (name_or_email or "").strip().lower()
    cleaned = _clean_name_input(name_or_email)
    if not raw:
        return None

    panel_users = _panel_interviewer_users(panel)
    if not panel_users:
        return None

    if "@" in raw:
        email_matches = [
            user for user in panel_users
            if (user.get("email") or "").strip().lower() == raw
        ]
        if len(email_matches) == 1:
            return email_matches[0]
        if len(email_matches) > 1:
            return "ambiguous"

    local_matches = [
        user for user in panel_users
        if (user.get("email") or "").strip().lower().split("@")[0] == raw
    ]
    if len(local_matches) == 1:
        return local_matches[0]
    if len(local_matches) > 1:
        return "ambiguous"

    if len(raw) >= 3:
        local_substring_matches = [
            user for user in panel_users
            if raw in (user.get("email") or "").strip().lower().split("@")[0]
        ]
        if len(local_substring_matches) == 1:
            return local_substring_matches[0]
        if len(local_substring_matches) > 1:
            return "ambiguous"

    for candidate in (cleaned, _normalize_person_name(name_or_email)):
        if not candidate:
            continue
        exact_name_matches = [
            user for user in panel_users
            if _normalize_person_name(user.get("name")) == candidate
        ]
        if len(exact_name_matches) == 1:
            return exact_name_matches[0]
        if len(exact_name_matches) > 1:
            return "ambiguous"

        partial_matches = [
            user for user in panel_users
            if candidate in _normalize_person_name(user.get("name"))
            or _normalize_person_name(user.get("name")) in candidate
        ]
        if len(partial_matches) == 1:
            return partial_matches[0]
        if len(partial_matches) > 1:
            return "ambiguous"

        first = candidate.split()[0]
        if first:
            first_matches = [
                user for user in panel_users
                if (_normalize_person_name(user.get("name")).split() or [""])[0] == first
            ]
            if len(first_matches) == 1:
                return first_matches[0]
            if len(first_matches) > 1:
                return "ambiguous"

    return None


def _panel_interviewer_labels(panel):
    return [
        f"{user.get('name') or 'Unknown'} ({user.get('email') or 'no email'})"
        for user in _panel_interviewer_users(panel)
    ]


def _interviewer_panels_for_dashboard():
    panels = {}
    for job_id, panel in INTERVIEWER_PANELS.items():
        job = JOBS.get(job_id, {})
        interviewers = []
        for user_id in panel.get("interviewer_ids") or []:
            interviewer = USERS.get(user_id)
            if not interviewer or interviewer.get("role") != "interviewer":
                continue
            interviewers.append({
                "id": user_id,
                "name": interviewer.get("name"),
                "email": interviewer.get("email"),
                "activated": True,
            })
        link = None
        access_code = None
        expires_display = ""
        if panel.get("token") and not _panel_link_expired(panel):
            link = f"{request.url_root.rstrip('/')}/interviewer/join/{panel['token']}"
            access_code = panel.get("access_code")
            try:
                expires_display = datetime.fromisoformat(panel["expires"]).strftime("%Y-%m-%d %H:%M")
            except Exception:
                expires_display = panel.get("expires") or ""
        panels[job_id] = {
            "job_id": job_id,
            "job_title": job.get("title", "Unknown Job"),
            "interviewers": interviewers,
            "link": link,
            "access_code": access_code,
            "expires": panel.get("expires"),
            "expires_display": expires_display,
            "has_active_link": bool(link),
            "interviewer_count": len(interviewers),
        }
    return panels


@app.route("/hr/interviewers/add", methods=["POST"])
def hr_add_interviewer():
    user, err = _require_hr_admin()
    if err:
        return err

    payload = request.get_json(silent=True) or request.form
    name = (payload.get("name") or "").strip()
    email = (payload.get("email") or "").strip().lower()
    job_id = (payload.get("job_id") or "").strip()

    if not name or not email or not job_id:
        return jsonify({"ok": False, "error": "Name, email, and job position are required"}), 400
    if not JOBS.get(job_id):
        return jsonify({"ok": False, "error": "Job position not found"}), 404

    for existing in USERS.values():
        if existing.get("email", "").lower() == email:
            return jsonify({"ok": False, "error": "Email already exists"}), 400

    import uuid
    new_user_id = f"user-{str(uuid.uuid4())[:8]}"
    new_user = {
        "id": new_user_id,
        "email": email,
        "name": name,
        "password_hash": generate_password_hash(secrets.token_urlsafe(16)),
        "role": "interviewer",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "can_screen_interviews": False,
        "can_score_interviews": True,
        "interviewer_token": None,
        "interviewer_token_expires": None,
        "interviewer_token_used": False,
        "interviewer_job_id": job_id,
        "password_reset_required": False,
        "first_login": False,
    }
    USERS[new_user_id] = new_user

    panel = _get_interviewer_panel(job_id, create=True)
    if new_user_id not in panel["interviewer_ids"]:
        panel["interviewer_ids"].append(new_user_id)
    panel["updated_at"] = datetime.now(timezone.utc).isoformat()

    save_users()
    save_interviewer_panels()
    log_activity(user["id"], "add_interviewer", f"Added interviewer {email} to job {job_id}", "success")
    return jsonify({
        "ok": True,
        "interviewer": {
            "id": new_user_id,
            "name": name,
            "email": email,
            "activated": False,
        },
        "panel": _interviewer_panels_for_dashboard().get(job_id),
    })


@app.route("/hr/interviewers/generate-link", methods=["POST"])
def hr_generate_interviewer_link():
    user, err = _require_hr_admin()
    if err:
        return err

    payload = request.get_json(silent=True) or request.form
    job_id = (payload.get("job_id") or "").strip()
    if not job_id:
        return jsonify({"ok": False, "error": "Job position is required"}), 400

    panel = _get_interviewer_panel(job_id)
    if not panel:
        return jsonify({"ok": False, "error": "Add at least one interviewer before generating a link"}), 400
    _sync_panel_membership(panel)
    if not panel.get("interviewer_ids"):
        return jsonify({"ok": False, "error": "Add at least one interviewer before generating a link"}), 400

    import uuid
    panel["token"] = str(uuid.uuid4())
    panel["access_code"] = _generate_panel_access_code()
    panel["expires"] = (datetime.now(timezone.utc) + timedelta(hours=24)).isoformat()
    panel["updated_at"] = datetime.now(timezone.utc).isoformat()
    save_interviewer_panels()

    link = f"{request.url_root.rstrip('/')}/interviewer/join/{panel['token']}"
    log_activity(user["id"], "generate_interviewer_link", f"Generated shared interviewer link for job {job_id}", "success")
    return jsonify({
        "ok": True,
        "link": link,
        "access_code": panel["access_code"],
        "expires": panel["expires"],
        "panel": _interviewer_panels_for_dashboard().get(job_id),
    })


@app.route("/hr/interviewers/create", methods=["POST"])
def hr_create_interviewer():
    """Backward-compatible alias: add interviewer only (no link)."""
    return hr_add_interviewer()


@app.route("/interviewer/login/<token>")
def interviewer_token_login(token):
    return redirect(url_for("interviewer_join", token=token))


@app.route("/interviewer/join/<token>", methods=["GET", "POST"])
def interviewer_join(token):
    panel = _get_panel_by_token(token)
    if not panel:
        flash("Invalid interviewer invitation link.", "danger")
        return redirect(url_for("login"))
    _sync_panel_membership(panel)
    if _panel_link_expired(panel):
        flash("This invitation link has expired (valid for 24 hours). Ask HR for a new link.", "danger")
        return redirect(url_for("login"))

    job = JOBS.get(panel.get("job_id"), {})
    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        access_code = (request.form.get("access_code") or "").strip()
        match = _find_panel_interviewer(panel, name)
        if match == "ambiguous":
            flash(
                "Multiple interviewers match that name. Use your full registered name or email address.",
                "danger",
            )
            return render_template(
                "interviewer_join.html",
                token=token,
                job_title=job.get("title", "Interview Panel"),
                panel_interviewers=_panel_interviewer_labels(panel),
            )
        if not match:
            registered = _panel_interviewer_labels(panel)
            if registered:
                flash(
                    "Name or email not found on this interview panel. "
                    f"Use exactly as HR registered you. On this panel: {'; '.join(registered)}",
                    "danger",
                )
            else:
                flash(
                    "No interviewers are registered on this panel yet. Ask HR to add you, then generate a new link.",
                    "danger",
                )
            return render_template(
                "interviewer_join.html",
                token=token,
                job_title=job.get("title", "Interview Panel"),
                panel_interviewers=registered,
            )

        panel_code = _normalize_access_code(panel.get("access_code"))
        entered_code = _normalize_access_code(access_code)
        if not panel_code or entered_code != panel_code:
            flash(
                "Incorrect access code. Use the 8-digit code HR shared when the link was generated (no spaces).",
                "danger",
            )
            return render_template(
                "interviewer_join.html",
                token=token,
                job_title=job.get("title", "Interview Panel"),
                panel_interviewers=_panel_interviewer_labels(panel),
            )

        session["user_id"] = match["id"]
        match["last_login"] = datetime.now(timezone.utc).isoformat()
        match["password_reset_required"] = False
        save_users()
        log_activity(match["id"], "interviewer_join", f"Interviewer {match.get('name')} joined via shared link", "success")
        flash("Welcome to the interviewer portal.", "success")
        return redirect(url_for("interviewer_dashboard"))

    return render_template(
        "interviewer_join.html",
        token=token,
        job_title=job.get("title", "Interview Panel"),
        panel_interviewers=_panel_interviewer_labels(panel),
    )


@app.route("/interviewer/set-password", methods=["GET", "POST"])
@app.route("/interviewer/reset-password/<token>", methods=["GET", "POST"])
def interviewer_reset_password_route(token=None):
    """Legacy route — interviewers now use a shared panel access code."""
    user = get_current_user()
    if user and user.get("role") == "interviewer":
        return redirect(url_for("interviewer_dashboard"))
    flash("Sign in with your invitation link and the access code from HR.", "info")
    if token:
        return redirect(url_for("interviewer_join", token=token))
    return redirect(url_for("login"))


@app.route("/interviewer/dashboard")
def interviewer_dashboard():
    user = get_current_user()
    if not user or user.get("role") != "interviewer":
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    job_id = user.get("interviewer_job_id")
    job = JOBS.get(job_id) if job_id else None
    panel = INTERVIEWER_PANELS.get(job_id, {}) if job_id else {}
    active_application_id = panel.get("active_application_id")

    candidates = []
    waiting_for_hr = False
    if job_id and active_application_id:
        app_data = APPLICATIONS.get(active_application_id)
        if app_data and app_data.get("job_id") == job_id and app_data.get("interview_status") == "in_progress":
            candidates.append(_enrich_application(app_data))
        else:
            waiting_for_hr = True
    elif job_id:
        waiting_for_hr = True

    return render_template(
        "interviewer_dashboard.html",
        job=job,
        candidates=candidates,
        waiting_for_hr=waiting_for_hr,
        active_changed_at=panel.get("active_changed_at"),
    )


@app.route("/interviewer/api/active-candidate")
def interviewer_active_candidate_api():
    user = get_current_user()
    if not user or user.get("role") != "interviewer":
        return jsonify({"ok": False, "error": "Unauthorized"}), 401

    job_id = user.get("interviewer_job_id")
    if not job_id:
        return jsonify({"ok": True, "active_application_id": None, "active_changed_at": None, "candidate": None})

    panel = INTERVIEWER_PANELS.get(job_id, {})
    active_application_id = panel.get("active_application_id")
    candidate = None
    if active_application_id:
        app_data = APPLICATIONS.get(active_application_id)
        if app_data and app_data.get("job_id") == job_id and app_data.get("interview_status") == "in_progress":
            candidate = _json_safe(_enrich_application(app_data))

    return jsonify({
        "ok": True,
        "active_application_id": active_application_id,
        "active_changed_at": panel.get("active_changed_at"),
        "candidate": candidate,
    })


@app.route("/hr/applications/<application_id>/data", methods=["GET"])
def hr_application_data(application_id):
    user = get_current_user()
    if not user or user.get("role") not in HR_ROLES:
        return jsonify({"ok": False, "error": "Unauthorized"}), 401
    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404
    return jsonify({"ok": True, "application": _json_safe(_enrich_application(app_data))})


@app.route("/hr/applications/<application_id>/resume")
def hr_application_resume(application_id):
    user, err = _require_hr()
    if err:
        return err
    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        return "Application not found", 404
    path = _resolve_upload_path(app_data.get("resume_path"))
    if not path:
        return "CV file not found", 404
    from flask import send_from_directory
    as_attachment = request.args.get("download") == "1"
    return send_from_directory(str(path.parent), path.name, as_attachment=as_attachment)


@app.route("/hr/applications/<application_id>/certificate/<int:cert_index>")
def hr_application_certificate(application_id, cert_index):
    user, err = _require_hr()
    if err:
        return err
    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        return "Application not found", 404
    cert_paths = app_data.get("certification_paths") or []
    if cert_index < 0 or cert_index >= len(cert_paths):
        return "Certificate not found", 404
    path = _resolve_upload_path(cert_paths[cert_index])
    if not path:
        return "Certificate file not found", 404
    from flask import send_from_directory
    as_attachment = request.args.get("download") == "1"
    return send_from_directory(str(path.parent), path.name, as_attachment=as_attachment)


@app.route("/hr/applications/<application_id>")
def hr_application_detail(application_id):
    user, err = _require_hr()
    if err:
        return err
    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        flash("Application not found.", "danger")
        return redirect(url_for("hr_screening"))
    application = _wrap_app(_enrich_application(app_data))
    job = _wrap_job(JOBS.get(app_data.get("job_id", "")) or {})
    return render_template("application_detail.html", application=application, job=job)


@app.route("/hr/applications/<application_id>/action", methods=["POST"])
def hr_application_action(application_id):
    user, err = _require_hr()
    if err:
        return err
    action = request.form.get("action", "")
    if application_id in APPLICATIONS:
        APPLICATIONS[application_id]["status"] = action
        try:
            with open(DATA_DIR / "applications.json", "w") as f:
                json.dump(APPLICATIONS, f, indent=2, default=str)
        except Exception as e:
            print(f"Warning: Could not save applications: {e}")
        flash(f"Application marked as {action}.", "success")
    return redirect(url_for("hr_application_detail", application_id=application_id))


@app.route("/hr/applications/<application_id>/screen", methods=["POST"])
def hr_application_screen(application_id):
    user, err = _require_hr()
    if err:
        return err

    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404

    payload = request.get_json(silent=True) or request.form
    raw_status = (payload.get("status") or "pending").strip().lower()
    status_map = {
        "shortlisted": "shortlisted",
        "pending review": "pending",
        "pending": "pending",
        "rejected": "rejected",
        "interview ready": "interview",
        "interview": "interview",
    }
    app_data["status"] = status_map.get(raw_status, raw_status or "pending")
    if app_data["status"] == "shortlisted":
        app_data["interview_status"] = "pending_scheduling"
    app_data["screening_rating"] = payload.get("rating")
    app_data["screening_notes"] = (payload.get("notes") or "").strip()
    app_data["screened_at"] = datetime.now(timezone.utc).isoformat()
    save_applications()
    log_activity(user["id"], "screen_candidate", f"Screened {application_id}", "success")
    return jsonify({"ok": True, "status": app_data["status"]})


@app.route("/hr/applications/<application_id>/reject", methods=["POST"])
def hr_application_reject(application_id):
    user, err = _require_hr()
    if err:
        return err

    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404

    payload = request.get_json(silent=True) or request.form
    send_email = _to_bool(payload.get("send_email", False))
    app_data["status"] = "rejected"
    app_data["rejection_reason"] = (payload.get("reason") or "").strip()
    app_data["rejection_message"] = (payload.get("message") or "").strip()
    app_data["rejected_at"] = datetime.now(timezone.utc).isoformat()

    email_result = None
    if send_email and app_data.get("email"):
        job = JOBS.get(app_data.get("job_id", ""), {})
        candidate_name = _candidate_full_name(app_data)
        format_values = {
            "candidate_name": candidate_name,
            "job_title": job.get("title", "the position"),
            "application_status": (app_data.get("status") or "rejected").upper(),
        }
        subject = f"Application Update - {candidate_name or 'Candidate'}"
        raw_body = app_data["rejection_message"] or "Thank you for your application. We will not move forward at this time."
        body = _personalize_message(raw_body, format_values)
        sent, error = send_email_message(
            app_data.get("email"), subject, body,
            sender_id=user["id"], application_id=application_id, email_type="rejection",
            candidate_name=candidate_name,
        )
        email_result = {"sent": sent, "error": error}
        app_data["rejection_email_sent"] = sent

    save_applications()
    log_activity(user["id"], "reject_candidate", f"Rejected {application_id}", "success")
    return jsonify({"ok": True, "email": email_result})


@app.route("/hr/interviews/schedule", methods=["POST"])
def hr_schedule_interview():
    user, err = _require_hr()
    if err:
        return err

    payload = request.get_json(silent=True) or request.form
    application_id = (payload.get("application_id") or "").strip()
    app_data = APPLICATIONS.get(application_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404

    interview_date = (payload.get("interview_date") or "").strip()
    interview_time = (payload.get("interview_time") or "").strip()
    if not interview_date or not interview_time:
        return jsonify({"ok": False, "error": "Interview date and time are required"}), 400

    app_data["status"] = "interview"
    app_data["interview_date"] = interview_date
    app_data["interview_time"] = interview_time
    app_data["interview_type"] = (payload.get("interview_type") or "Interview").strip()
    app_data["interviewer"] = (payload.get("interviewer") or "").strip()
    app_data["interview_location"] = (payload.get("interview_location") or "").strip()
    app_data["interview_message"] = (payload.get("message") or "").strip()
    app_data["interview_status"] = "scheduled"
    app_data["interview_removed"] = False
    app_data["interview_removed_reason"] = ""
    app_data["interview_scheduled_at"] = datetime.now(timezone.utc).isoformat()

    email_result = None
    send_email = _to_bool(payload.get("send_email", True))
    if send_email:
        if not app_data.get("email"):
            email_result = {"sent": False, "error": "Candidate has no email address"}
            app_data["interview_email_sent"] = False
            app_data["interview_email_error"] = email_result["error"]
        else:
            job = JOBS.get(app_data.get("job_id", ""), {})
            template = app_data["interview_message"] or SETTINGS.get(
                "interview_default_message", DEFAULT_SETTINGS["interview_default_message"]
            )
            format_values = {
                "candidate_name": _candidate_full_name(app_data),
                "job_title": job.get("title", "the position"),
                "interview_date": interview_date,
                "interview_time": interview_time,
                "interview_type": app_data.get("interview_type", "Interview"),
                "interview_location": app_data.get("interview_location", "To be confirmed"),
            }
            body = _personalize_message(template, format_values)
            if interview_date not in body or interview_time not in body:
                body = (
                    f"{body}\n\nInterview schedule:\n"
                    f"Date: {interview_date}\n"
                    f"Time: {interview_time}\n"
                    f"Type: {format_values['interview_type']}"
                )
            subject = f"Interview Invitation - {job.get('title', 'Application')}"
            sent, error = send_email_message(
                app_data.get("email"), subject, body,
                sender_id=user["id"], application_id=application_id, email_type="interview",
                candidate_name=format_values["candidate_name"],
                footer_location=format_values["interview_location"],
            )
            email_result = {"sent": sent, "error": error}
            app_data["interview_email_sent"] = sent
            app_data["interview_email_error"] = error

    save_applications()
    log_activity(user["id"], "schedule_interview", f"Scheduled interview for {application_id}", "success")
    return jsonify({"ok": True, "email": email_result, "application": _enrich_application(app_data)})


# ─── INTERVIEW ACTION & STAGE 2 SCREENING ENDPOINTS ───────────────────

def _require_scoring_permission(application_id=None):
    user = get_current_user()
    if not user:
        return None, (jsonify({"ok": False, "error": "Please log in."}), 401)
    if user.get("role") == "hr":
        return user, None
    if user.get("role") == "hr_staff":
        return user, None
    if user.get("role") == "interviewer" and application_id:
        app = APPLICATIONS.get(application_id)
        if app and app.get("job_id") == user.get("interviewer_job_id"):
            return user, None
    return None, (jsonify({"ok": False, "error": "You do not have permission to score this candidate"}), 403)


def _require_screening_permission():
    user = get_current_user()
    if not user:
        return None, (jsonify({"ok": False, "error": "Unauthorized"}), 401)
    if user.get("role") == "hr":
        return user, None
    if user.get("role") == "hr_staff" and user.get("can_screen_interviews"):
        return user, None
    return None, (jsonify({"ok": False, "error": "You do not have permission to screen interviews"}), 403)


@app.route("/hr/interviews/<app_id>/start", methods=["POST"])
def hr_start_interview(app_id):
    user, err = _require_hr_api()
    if err:
        return err
    app_data = APPLICATIONS.get(app_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404
    if app_data.get("interview_removed"):
        return jsonify({"ok": False, "error": "This candidate was removed from the interview queue."}), 400
    if app_data.get("interview_status") == "completed":
        return jsonify({"ok": False, "error": "This interview is closed and cannot be started again."}), 400
    if app_data.get("interview_status") == "cancelled":
        return jsonify({"ok": False, "error": "This candidate was cancelled as a no-show."}), 400
    if not (app_data.get("interview_date") or "").strip():
        return jsonify({
            "ok": False,
            "error": "Schedule this interview first (date and time) before clicking Start.",
        }), 400
    job_id = app_data.get("job_id")
    if not job_id:
        return jsonify({"ok": False, "error": "Candidate is not linked to a job"}), 400
    try:
        _set_active_interview(job_id, app_id)
        app_data["interview_status"] = "in_progress"
        save_applications()
        save_interviewer_panels()
    except Exception as ex:
        print(f"Error starting interview for {app_id}: {ex}")
        return jsonify({"ok": False, "error": "Could not start interview. Please try again."}), 500
    log_activity(user["id"], "start_interview", f"Started interview for {app_id}", "success")
    return jsonify({"ok": True, "application": _enrich_application(app_data)})


@app.route("/hr/interviews/<app_id>/end", methods=["POST"])
def hr_end_interview(app_id):
    user, err = _require_hr_api()
    if err:
        return err
    app_data = APPLICATIONS.get(app_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404
    _sync_interview_status_after_score(app_data)
    if app_data.get("interview_status") not in ("awaiting_screening", "completed"):
        app_data["interview_status"] = "in_progress"
    save_applications()
    log_activity(user["id"], "end_interview", f"Ended interview session for {app_id}", "success")
    return jsonify({"ok": True, "application": _enrich_application(app_data)})


@app.route("/hr/interviews/<app_id>/cancel", methods=["POST"])
def hr_cancel_interview(app_id):
    user, err = _require_hr_api()
    if err:
        return err
    app_data = APPLICATIONS.get(app_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404
    app_data["interview_status"] = "cancelled"
    app_data["status"] = "rejected"
    save_applications()
    log_activity(user["id"], "cancel_interview", f"Cancelled interview for {app_id} (No Show)", "success")
    return jsonify({"ok": True, "application": _enrich_application(app_data)})


@app.route("/hr/interviews/<app_id>/remove", methods=["POST"])
def hr_remove_interview_queue(app_id):
    user, err = _require_hr_api()
    if err:
        return err
    app_data = APPLICATIONS.get(app_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404
    payload = request.get_json(silent=True) or request.form
    reason = (payload.get("reason") or "").strip()
    if not reason:
        return jsonify({"ok": False, "error": "Reason is required to remove from queue"}), 400
    app_data["interview_removed"] = True
    app_data["interview_removed_reason"] = reason
    app_data["interview_status"] = None
    job_id = app_data.get("job_id")
    if job_id:
        panel = _get_interviewer_panel(job_id, create=False)
        if panel and panel.get("active_application_id") == app_id:
            panel["active_application_id"] = None
            panel["active_changed_at"] = datetime.now(timezone.utc).isoformat()
            save_interviewer_panels()
    save_applications()
    log_activity(user["id"], "remove_interview_queue", f"Removed {app_id} from interview queue: {reason}", "success")
    return jsonify({"ok": True, "application": _enrich_application(app_data)})


@app.route("/hr/interviews/<app_id>/score", methods=["POST"])
def hr_score_interview(app_id):
    user, err = _require_scoring_permission(app_id)
    if err:
        return err
    app_data = APPLICATIONS.get(app_id)
    if not app_data:
        return jsonify({"ok": False, "error": "Application not found"}), 404
    if app_data.get("interview_status") == "completed":
        return jsonify({"ok": False, "error": "Interview is completed. Marks cannot be changed."}), 400
    if user.get("role") == "interviewer" and not _is_active_interview_candidate(app_data):
        return jsonify({
            "ok": False,
            "error": "This candidate is not active right now. Wait for HR to start their interview.",
        }), 400

    payload = request.get_json(silent=True) or request.form

    scores = {}
    total_score = 0
    for crit in INTERVIEW_SCORE_CRITERIA:
        val_raw = payload.get(crit, 0)
        try:
            val = int(val_raw)
        except (ValueError, TypeError):
            val = 0
        val = max(0, min(20, val))
        scores[crit] = val
        total_score += val

    percentage = (total_score / 140.0) * 100.0
    _upsert_interview_score_submission(app_data, user, scores, percentage)
    _sync_interview_status_after_score(app_data)

    save_applications()
    log_activity(user["id"], "score_interview", f"Scored interview for {app_id} ({app_data['interview_total']}%)", "success")
    return jsonify({"ok": True, "application": _enrich_application(app_data)})


@app.route("/hr/interviews/screen", methods=["POST"])
def hr_screen_interviews():
    user, err = _require_screening_permission()
    if err:
        return err
    payload = request.get_json(silent=True) or request.form
    job_id = (payload.get("job_id") or "").strip()
    if not job_id:
        return jsonify({"ok": False, "error": "Job ID is required"}), 400

    job = JOBS.get(job_id)
    if not job:
        return jsonify({"ok": False, "error": "Job not found"}), 404

    queue = _interview_queue_for_job(job_id)
    if not queue:
        return jsonify({"ok": False, "error": "No interview candidates found for this job."}), 400

    not_started = [
        _candidate_full_name(app)
        for app in queue
        if app.get("interview_status") in (None, "scheduled")
    ]
    if not_started:
        return jsonify({
            "ok": False,
            "error": "Cannot screen yet. Start and finish interviews for all candidates first: " + ", ".join(not_started),
        }), 400

    incomplete = []
    candidates_list = []
    for app in queue:
        if not _all_required_scorers_submitted(app):
            pending = _pending_scorer_names(app)
            incomplete.append(f"{_candidate_full_name(app)} (waiting: {', '.join(pending)})")
            continue
        app["interview_status"] = "completed"
        app["interview_screened_at"] = datetime.now(timezone.utc).isoformat()
        candidates_list.append(_enrich_application(app))

    if incomplete:
        return jsonify({
            "ok": False,
            "error": "Cannot screen yet. All panel interviewers must submit marks for every candidate. Pending: " + "; ".join(incomplete),
        }), 400

    if not candidates_list:
        return jsonify({
            "ok": False,
            "error": "No interviewed candidates are ready for screening yet.",
        }), 400

    save_applications()
    candidates_list.sort(key=lambda x: x.get("interview_average", x.get("interview_total", 0)), reverse=True)
    log_activity(user["id"], "screen_interviews", f"Screened {len(candidates_list)} interview(s) for job {job_id}", "success")
    return jsonify({"ok": True, "candidates": candidates_list})


@app.route("/hr/interviews/select-best", methods=["POST"])
def hr_select_best_candidates():
    user, err = _require_screening_permission()
    if err:
        return err
    payload = request.get_json(silent=True) or request.form
    job_id = (payload.get("job_id") or "").strip()
    try:
        positions_needed = int(payload.get("positions_needed", 1))
    except (ValueError, TypeError):
        positions_needed = 1

    if not job_id:
        return jsonify({"ok": False, "error": "Job ID is required"}), 400

    job = JOBS.get(job_id)
    if not job:
        return jsonify({"ok": False, "error": "Job not found"}), 404

    # Find completed non-removed applications for this job
    candidates_with_scores = []
    for app in APPLICATIONS.values():
        if app.get("job_id") == job_id and app.get("interview_status") == "completed" and not app.get("interview_removed"):
            candidates_with_scores.append(app)

    if not candidates_with_scores:
        return jsonify({"ok": False, "error": "No completed interviews found for this job"}), 400

    def _candidate_average(app_item):
        return _interview_score_average(app_item)

    candidates_with_scores.sort(key=_candidate_average, reverse=True)

    from collections import defaultdict
    groups = defaultdict(list)
    for c in candidates_with_scores:
        groups[_candidate_average(c)].append(c)

    sorted_scores = sorted(groups.keys(), reverse=True)
    selected_candidates = []
    remaining_slots = positions_needed
    import random

    for score in sorted_scores:
        group = groups[score]
        if len(group) <= remaining_slots:
            selected_candidates.extend(group)
            remaining_slots -= len(group)
        else:
            # Tiebreak needed
            chosen = random.sample(group, remaining_slots)
            for ch in chosen:
                ch["selection_method"] = "random_tiebreak"
            selected_candidates.extend(chosen)
            remaining_slots = 0
        if remaining_slots <= 0:
            break

    # Update selected candidates
    now_iso = datetime.now(timezone.utc).isoformat()
    for sc in selected_candidates:
        if sc.get("selection_method") != "random_tiebreak":
            sc["selection_method"] = "ranked"
        sc["selected"] = True
        sc["status"] = "selected"
        sc["selected_at"] = now_iso

    save_applications()
    log_activity(user["id"], "select_candidates", f"Selected {len(selected_candidates)} candidates for {job.get('title')}", "success")
    return jsonify({"ok": True, "selected": [_enrich_application(sc) for sc in selected_candidates]})


@app.route("/hr/selected/send-email", methods=["POST"])
def hr_send_selected_email():
    user, err = _require_hr()
    if err:
        return err
    payload = request.get_json(silent=True) or request.form
    job_id = (payload.get("job_id") or "").strip()
    subject = (payload.get("subject") or "").strip()
    message = (payload.get("message") or "").strip()
    start_date = (payload.get("start_date") or "").strip()
    start_time = (payload.get("start_time") or "").strip()
    location = (payload.get("location") or "").strip()

    if not job_id or not subject or not message:
        return jsonify({"ok": False, "error": "Job ID, subject, and message are required"}), 400

    job = JOBS.get(job_id)
    if not job:
        return jsonify({"ok": False, "error": "Job not found"}), 404

    # Find all selected candidates for this job
    selected_apps = [app for app in APPLICATIONS.values() if app.get("job_id") == job_id and app.get("selected")]
    if not selected_apps:
        return jsonify({"ok": False, "error": "No selected candidates found for this job"}), 400

    sent_count = 0
    failed_count = 0
    errors = []

    for app_data in selected_apps:
        candidate_name = _candidate_full_name(app_data)
        format_values = {
            "candidate_name": candidate_name,
            "job_title": job.get("title", "the position"),
            "start_date": start_date or "To be confirmed",
            "start_time": start_time or "To be confirmed",
            "location": location or "To be confirmed",
        }
        body = _personalize_message(message, format_values)

        sent, error = send_email_message(
            app_data.get("email"), subject, body,
            sender_id=user["id"], application_id=app_data.get("id"), email_type="selection",
            candidate_name=candidate_name,
        )
        if sent:
            sent_count += 1
            app_data["selected_email_sent"] = True
        else:
            failed_count += 1
            errors.append(f"Failed for {candidate_name}: {error}")

    save_applications()
    log_activity(user["id"], "send_selection_emails", f"Sent selection emails for {job.get('title')}: {sent_count} success, {failed_count} failed", "success")
    return jsonify({"ok": True, "sent": sent_count, "failed": failed_count, "errors": errors})


@app.route("/hr/bulk-email", methods=["POST"])
def hr_bulk_email():
    """Send bulk emails to applicants filtered by status"""
    user, err = _require_hr()
    if err:
        return err

    payload = request.get_json(silent=True) or request.form
    subject = (payload.get("subject") or "").strip()
    message = (payload.get("message") or "").strip()
    status_filter = (payload.get("filter") or "all").strip()
    group_by_job = _to_bool(payload.get("group_by_job", True))

    if not subject or not message:
        return jsonify({"ok": False, "error": "Subject and message are required"}), 400

    filtered_apps = []
    for app_data in APPLICATIONS.values():
        if status_filter == "all" or app_data.get("status") == status_filter:
            filtered_apps.append(app_data)

    if not filtered_apps:
        return jsonify({"ok": False, "error": "No applicants found with selected filter"}), 400

    if group_by_job:
        from collections import defaultdict
        grouped = defaultdict(list)
        for app_data in filtered_apps:
            job_title = JOBS.get(app_data.get("job_id", ""), {}).get("title", "General Application")
            grouped[job_title].append(app_data)
        send_batches = list(grouped.items())
    else:
        send_batches = [("All Applicants", filtered_apps)]

    sent_count = 0
    failed_count = 0
    job_summaries = []

    for job_title, apps_in_job in send_batches:
        job_sent = 0
        for app_data in apps_in_job:
            email = app_data.get("email")
            if not email:
                failed_count += 1
                continue

            format_values = {
                "candidate_name": _candidate_full_name(app_data),
                "job_title": JOBS.get(app_data.get("job_id", ""), {}).get("title", job_title),
                "application_status": (app_data.get("status") or "pending").upper(),
            }

            body = _personalize_message(message, format_values)
            personalized_subject = _personalize_message(subject, format_values)
            if group_by_job and "(job title)" not in subject.lower() and "{job_title}" not in subject:
                personalized_subject = f"[{format_values['job_title']}] {personalized_subject}"

            sent, error = send_email_message(
                email, personalized_subject, body,
                sender_id=user["id"], application_id=app_data.get("id"), email_type="bulk",
                candidate_name=format_values["candidate_name"],
            )
            if sent:
                sent_count += 1
                job_sent += 1
            else:
                failed_count += 1
        if job_sent:
            job_summaries.append(f"{job_title}: {job_sent}")

    save_applications()
    log_activity(
        user["id"], "bulk_email",
        f"Sent bulk email to {sent_count} applicants (filter: {status_filter}, grouped: {group_by_job})",
        "success",
    )
    return jsonify({
        "ok": True,
        "count": len(filtered_apps),
        "sent": sent_count,
        "failed": failed_count,
        "grouped_by_job": group_by_job,
        "job_summaries": job_summaries,
    })


@app.route("/applicants/<application_id>")
def applicant_page(application_id):
    user = get_current_user()
    if user and user.get("role") in HR_ROLES:
        return redirect(url_for("hr_dashboard"))
        flash("Application not found.", "danger")
        return redirect(url_for("index"))


@app.route("/hr/report")
def hr_report_summary():
    user, err = _require_hr()
    if err:
        return err
    status_filter = (request.args.get("status") or "all").strip().lower()
    report_title = (request.args.get("report_title") or "Applicants Report").strip()
    report_notes = (request.args.get("notes") or "").strip()
    requested_filename = (request.args.get("filename") or "report").strip()
    fmt = _normalize_report_format(request.args.get("format"))

    spec = _build_summary_report_spec(user, status_filter=status_filter, report_title=report_title, report_notes=report_notes)
    return _send_report_download(spec, requested_filename, fmt)


@app.route("/hr/audit-log")
def hr_audit_log():
    flash("Audit trail is only available to audit users.", "warning")
    return redirect(url_for("hr_dashboard"))


@app.route("/hr/reports/cv-analysis")
@app.route("/hr/reports/cv-analysis.csv")
def hr_cv_analysis_report():
    """Generate AI-powered CV analysis report for all applications"""
    user, err = _require_hr()
    if err:
        return err
    fmt = _normalize_report_format(request.args.get("format"))
    requested_filename = (request.args.get("filename") or "cv-analysis-report").strip()
    spec = _build_cv_analysis_report_spec(user)
    return _send_report_download(spec, requested_filename, fmt)


# ─── Public job / apply routes ─────────────────────────────────────────────
@app.route("/jobs/<job_id>")
def job_detail(job_id):
    job_data = JOBS.get(job_id)
    if not job_data:
        flash("Job not found.", "danger")
        return redirect(url_for("index"))
    job = _wrap_job(_enrich_job(job_data))
    application_link = url_for("apply", job_id=job_id, _external=True)
    return render_template("job_detail.html", job=job, application_link=application_link)


@app.route("/jobs/<job_id>/apply", methods=["GET", "POST"])
def apply(job_id):
    job_data = JOBS.get(job_id)
    if not job_data:
        flash("Job not found.", "danger")
        return redirect(url_for("index"))
    source = (request.args.get("source") or "").strip().lower()
    show_hr_sidebar = source == "hr"
    education_options = ["High School", "Diploma", "Bachelor's Degree", "Master's Degree", "PhD"]
    job = _wrap_job(_enrich_job(job_data))
    if request.method == "POST":
        import uuid
        from werkzeug.utils import secure_filename

        form_email = request.form.get("email", "").strip()
        if _has_application_for_job(job_id, form_email):
            flash("You have already submitted an application for this position. Only one CV submission is allowed per job.", "warning")
            return redirect(url_for("apply", job_id=job_id))

        app_id = str(uuid.uuid4())
        resume_path = ""
        certification_paths = []
        upload_dir = Path(__file__).parent / "uploads"
        upload_dir.mkdir(exist_ok=True)

        resume_file = request.files.get("resume")
        if not resume_file or not resume_file.filename:
            flash("Please upload exactly one CV file.", "danger")
            return redirect(request.url)

        filename = secure_filename(resume_file.filename)
        save_path = upload_dir / f"{app_id}_cv_{filename}"
        resume_file.save(str(save_path))
        resume_path = str(save_path)

        cert_files = [f for f in request.files.getlist("certifications") if f and f.filename]
        if len(cert_files) > 100:
            flash("You can upload a maximum of 100 certificate files.", "danger")
            Path(resume_path).unlink(missing_ok=True)
            return redirect(request.url)

        for index, cert_file in enumerate(cert_files, start=1):
            cert_filename = secure_filename(cert_file.filename)
            cert_path = upload_dir / f"{app_id}_cert_{index}_{cert_filename}"
            cert_file.save(str(cert_path))
            certification_paths.append(str(cert_path))

        form_name = request.form.get("name", "").strip()
        form_surname = request.form.get("surname", "").strip()
        form_phone = request.form.get("phone", "").strip()
        form_education = request.form.get("highest_education", "").strip()

        try:
            from services.document_validation import validate_application_documents
            validation_errors = validate_application_documents(
                form_name,
                form_surname,
                form_email,
                form_phone,
                form_education,
                Path(resume_path),
                [Path(p) for p in certification_paths],
            )
        except Exception as exc:
            validation_errors = [f"Could not validate uploaded documents: {exc}"]

        if validation_errors:
            Path(resume_path).unlink(missing_ok=True)
            for cert_ref in certification_paths:
                Path(cert_ref).unlink(missing_ok=True)
            flash(validation_errors[0], "danger")
            return redirect(request.url)

        APPLICATIONS[app_id] = {
            "id": app_id,
            "job_id": job_id,
            "name": form_name,
            "surname": form_surname,
            "email": form_email,
            "phone": form_phone,
            "highest_education": form_education,
            "note": request.form.get("note", "").strip(),
            "resume_path": resume_path,
            "certification_paths": certification_paths,
            "status": "pending",
            "score": 0,
            "summary": "",
            "recommendation": "",
            "matched_skills": [],
            "missing_skills": [],
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
        }
        try:
            with open(DATA_DIR / "applications.json", "w") as f:
                json.dump(APPLICATIONS, f, indent=2, default=str)
        except Exception as e:
            print(f"Warning: Could not save application: {e}")
        return redirect(url_for("success_page", application_id=app_id))
    return render_template(
        "apply.html",
        job=job,
        education_options=education_options,
        show_hr_sidebar=show_hr_sidebar,
    )


@app.route("/success")
def success_page():
    application_id = request.args.get("application_id", "")
    app_data = APPLICATIONS.get(application_id) or {}
    application = _wrap_app(app_data)
    job = _wrap_job(JOBS.get(app_data.get("job_id", "")) or {})
    return render_template("success.html", application=application, job=job)


@app.route("/uploads/<filename>")
def uploaded_file(filename):
    from flask import send_from_directory
    from werkzeug.utils import secure_filename
    safe_name = secure_filename(filename)
    upload_dir = Path(__file__).parent / "uploads"
    target = upload_dir / safe_name
    if not target.exists():
        return "File not found", 404
    as_attachment = request.args.get("download") == "1"
    return send_from_directory(str(upload_dir), safe_name, as_attachment=as_attachment)


load_all_data()


if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("HR System - Flask Application")
    print("=" * 60)
    print("URL:  http://localhost:5000")
    print("HR:   hr@company.com / password123")
    print("Audit: audit@company.com / password123")
    print("=" * 60 + "\n")
    app.run(debug=True, host="0.0.0.0", port=5000, use_reloader=True)
