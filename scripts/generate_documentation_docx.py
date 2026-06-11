#!/usr/bin/env python
"""Generate PROJECT_DOCUMENTATION.docx from markdown source."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "PROJECT_DOCUMENTATION.docx"


def set_cell_shading(cell, fill_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="124E66"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_fill)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_document():
    doc = Document()

    # Title page
    title = doc.add_heading("HR Recruitment System", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Project Documentation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    org = doc.add_paragraph("Grain Marketing Board (GMB) HR Portal")
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.runs[0].font.color.rgb = RGBColor(18, 78, 102)
    doc.add_paragraph(
        "A Flask-based recruitment and applicant tracking system for posting jobs, "
        "screening candidates, scheduling interviews, and managing external interview panels."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Last updated: June 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of contents (manual)
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Overview",
        "2. Portals & User Roles",
        "3. Main Features",
        "4. Project Structure",
        "5. Where to Put Pictures & Media",
        "6. Installation & Running",
        "7. Configuration",
        "8. Default Logins",
        "9. Typical Workflows",
        "10. Data Storage",
        "11. API & Key Routes",
        "12. Branding & Colors",
        "13. Troubleshooting",
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph("This system helps HR teams manage the full hiring lifecycle:")
    add_table(
        doc,
        ["Stage", "What the system does"],
        [
            ("Post jobs", "HR creates job listings with title, description, department, and positions needed"),
            ("Apply", "Candidates browse open jobs and submit applications with CV and certificates"),
            ("Screen", "CVs are scored and ranked against job requirements"),
            ("Interview", "HR schedules interviews and sends email invitations"),
            ("Score", "HR staff and external interviewers evaluate candidates on criteria"),
            ("Select", "Best candidates are shortlisted and notified"),
        ],
    )
    doc.add_paragraph(
        "Tech stack: Python 3, Flask, Jinja2 templates, JSON file storage, "
        "optional PostgreSQL hooks, SMTP email, scikit-learn CV scoring."
    )

    # 2. Portals
    doc.add_heading("2. Portals & User Roles", level=1)

    doc.add_heading("2.1 Applicant Portal (Public)", level=2)
    add_bullets(doc, [
        "URL: http://localhost:5000/ or /portal",
        "Who: Job seekers (no login required)",
        "Pages: Job listings, job detail, application form, success page",
    ])

    doc.add_heading("2.2 HR Portal (Staff)", level=2)
    add_bullets(doc, [
        "URL: http://localhost:5000/login → /hr",
        "HR Manager (hr): full access, user management, settings",
        "HR Staff (hr_staff): screening/scoring based on permissions",
        "Sections: Dashboard, Jobs, Screening, Interviews, Interview Marks, Selected Candidates, Settings, Reports",
    ])

    doc.add_heading("2.3 Interviewer Portal (External)", level=2)
    add_bullets(doc, [
        "URL: Shared link from HR, e.g. /interviewer/join/<token>",
        "Who: External interview panel members",
        "Access: Full name + shared 8-digit panel access code",
        "Purpose: Score shortlisted candidates for assigned job",
    ])

    doc.add_heading("2.4 Audit Portal", level=2)
    add_bullets(doc, [
        "URL: /login → /audit",
        "Who: Audit managers and audit members",
        "Purpose: Review system activity and compliance",
    ])

    # 3. Features
    doc.add_heading("3. Main Features", level=1)
    doc.add_heading("Recruitment", level=2)
    add_bullets(doc, [
        "Create and manage job postings (positions_needed per job)",
        "Public careers page with search and filters",
        "One application per email per job",
        "CV upload (PDF, DOCX) with automatic scoring",
    ])
    doc.add_heading("Screening", level=2)
    add_bullets(doc, [
        "AI-assisted CV ranking (TF-IDF + skill matching)",
        "Manual screening with notes",
        "Reject candidates with optional email",
    ])
    doc.add_heading("Interviews", level=2)
    add_bullets(doc, [
        "Schedule interviews with date, time, type, location",
        "Personalized email invitations (HTML with company logo)",
        "Interview status: scheduled → in progress → awaiting screening → completed",
        "Interview Marks criteria: presentation, dressing, communication, confidence, technical knowledge, problem solving, attitude",
        "Completed only after all panel interviewers score and HR runs Screen Completed Interviews",
    ])
    doc.add_heading("External Interviewers", level=2)
    add_bullets(doc, [
        "Add multiple interviewers per job",
        "Generate one shared link + one access code (24-hour expiry)",
        "Interviewers sign in and submit scores; marks appear in HR Interview Marks",
    ])
    doc.add_heading("Communication & Settings", level=2)
    add_bullets(doc, [
        "SMTP email for interviews, rejections, bulk messages",
        "Email history log",
        "Company logo embedded in outgoing emails",
        "Company name, logo, welcome image, welcome video",
        "Mail server configuration and user management",
    ])

    # 4. Structure
    doc.add_heading("4. Project Structure", level=1)
    add_code(
        doc,
        "hr/\n"
        "├── app.py                 # Main Flask application\n"
        "├── requirements.txt       # Python dependencies\n"
        "├── data/                  # JSON database\n"
        "├── static/                # CSS, default logo\n"
        "├── uploads/               # CVs, branding files\n"
        "├── templates/             # HTML pages\n"
        "├── services/              # CV scoring, email, validation\n"
        "└── docs/images/           # Documentation screenshots",
    )

    # 5. Pictures
    doc.add_heading("5. Where to Put Pictures & Media", level=1)

    doc.add_heading("5.1 Company Logo", level=2)
    doc.add_paragraph("Best method: Upload in HR → Settings → Portal Branding.")
    add_table(
        doc,
        ["Use", "Where it appears"],
        [
            ("HR login page", "templates/login.html"),
            ("HR dashboard sidebar", "templates/hr_dashboard.html"),
            ("Interviewer portal", "interviewer_join.html, interviewer_dashboard.html"),
            ("Applicant portal header", "templates/index.html"),
            ("Outgoing emails", "Embedded as company_logo"),
        ],
    )
    add_bullets(doc, [
        "File saved to: uploads/branding/company_logo.png",
        "Fallback default: static/gmb-logo.svg",
        "Supported formats: PNG, JPG, SVG",
    ])

    doc.add_heading("5.2 Welcome Image & Video", level=2)
    add_bullets(doc, [
        "Welcome image: HR → Settings → Portal Branding (stored in data/settings.json)",
        "Alternative: static/images/welcome-banner.jpg",
        "Welcome video: uploads/branding/your-video.mp4 via Settings",
        "Use MP4 (H.264), keep under 50 MB where possible",
    ])

    doc.add_heading("5.3 Applicant CVs (Automatic)", level=2)
    add_code(doc, "uploads/<application-id>_cv_<filename>.pdf")
    doc.add_paragraph("Do not place these manually — the apply form handles uploads.")

    doc.add_heading("5.4 Documentation Screenshots", level=2)
    add_code(
        doc,
        "docs/images/\n"
        "├── 01-applicant-portal.png\n"
        "├── 02-hr-dashboard.png\n"
        "├── 03-interview-marks.png\n"
        "├── 04-interviewer-portal.png\n"
        "└── 05-settings-branding.png",
    )
    doc.add_paragraph("These are for reports and presentations only — not used by the running app.")

    doc.add_heading("5.5 Quick Reference", level=2)
    add_table(
        doc,
        ["What you want", "Where to put it", "How to activate"],
        [
            ("Company logo", "HR Settings or static/gmb-logo.svg", "Settings → Portal Branding"),
            ("Welcome banner", "HR Settings or static/images/", "Settings or edit index.html"),
            ("Welcome video", "uploads/branding/ via Settings", "Settings → Portal Branding"),
            ("Applicant CVs", "uploads/ (automatic)", "Candidate applies online"),
            ("CSS / styling", "static/style.css", "Edit file directly"),
            ("Doc screenshots", "docs/images/", "Manual — for reports only"),
        ],
    )

    # 6. Installation
    doc.add_heading("6. Installation & Running", level=1)
    doc.add_paragraph("Prerequisites: Python 3.11+, pip")
    doc.add_heading("Windows Steps", level=2)
    add_code(
        doc,
        "cd C:\\Users\\PC\\hr\n"
        "python -m venv venv\n"
        "venv\\Scripts\\activate\n"
        "pip install -r requirements.txt\n"
        "python app.py",
    )
    doc.add_paragraph("Open: http://localhost:5000")
    doc.add_paragraph("Health check: GET http://localhost:5000/health")

    # 7. Configuration
    doc.add_heading("7. Configuration", level=1)
    doc.add_paragraph("Create a .env file in the project root:")
    add_code(
        doc,
        "SECRET_KEY=your-secret-key-here\n"
        "MAIL_SERVER=smtp.gmail.com\n"
        "MAIL_PORT=587\n"
        "MAIL_USE_TLS=true\n"
        "MAIL_USERNAME=your-email@gmail.com\n"
        "MAIL_PASSWORD=your-app-password\n"
        "MAIL_DEFAULT_SENDER=your-email@gmail.com",
    )
    doc.add_paragraph(
        "For Gmail, use an App Password, not your normal login password. "
        "Most branding and mail options can also be changed from HR → Settings."
    )

    # 8. Logins
    doc.add_heading("8. Default Logins", level=1)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ("HR Manager", "hr@company.com", "password123"),
            ("Audit Manager", "audit@company.com", "password123"),
        ],
    )
    doc.add_paragraph(
        "Interviewer access: HR generates a shared link + 8-digit access code per job panel. "
        "Change default passwords before deploying to production."
    )

    # 9. Workflows
    doc.add_heading("9. Typical Workflows", level=1)
    doc.add_heading("HR: Full hiring flow", level=2)
    steps = [
        "Login at /login",
        "Post a job — HR Dashboard → Jobs (set positions needed)",
        "Wait for applications — candidates apply via /portal",
        "Screen candidates — Screening section, run CV scoring",
        "Schedule interviews — Interviews section, send email invite",
        "Add external interviewers — add names, generate shared link + access code",
        "Review marks — Interview Marks section",
        "Screen completed interviews when all evaluators have scored",
        "Select candidates — Selected Candidates section",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("Interviewer: Score candidates", level=2)
    for i, step in enumerate([
        "Open shared link from HR",
        "Enter full name (as registered) + panel access code",
        "Evaluate candidates and submit scores",
        "Scores appear in HR Interview Marks",
    ], 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("Applicant: Apply for a job", level=2)
    for i, step in enumerate([
        "Visit / or /portal",
        "Browse/search jobs",
        "Click Apply → fill form → upload CV",
        "Confirmation on success page",
    ], 1):
        doc.add_paragraph(f"{i}. {step}")

    # 10. Data
    doc.add_heading("10. Data Storage", level=1)
    add_table(
        doc,
        ["File", "Contents"],
        [
            ("users.json", "HR, audit, interviewer accounts"),
            ("jobs.json", "Job postings"),
            ("applications.json", "Candidate applications and interview data"),
            ("settings.json", "Company name, logo, mail config"),
            ("interviewer_panels.json", "Shared links, access codes, panel members"),
            ("activity_logs.json", "Audit trail of user actions"),
            ("email_history.json", "Sent email log"),
        ],
    )
    doc.add_paragraph("Backup tip: Copy the entire data/ and uploads/ folders regularly.")

    # 11. Routes
    doc.add_heading("11. API & Key Routes", level=1)
    add_table(
        doc,
        ["Route", "Description"],
        [
            ("/", "Applicant job listings"),
            ("/login", "HR / Audit login"),
            ("/hr", "HR dashboard"),
            ("/jobs/<id>/apply", "Application form"),
            ("/hr/jobs/post", "Create job (POST)"),
            ("/hr/interviews/schedule", "Schedule interview"),
            ("/hr/interviewers/generate-link", "Generate panel link + code"),
            ("/interviewer/join/<token>", "Interviewer sign-in"),
            ("/interviewer/dashboard", "Interviewer scoring portal"),
            ("/hr/settings/portal-branding", "Upload logo, image, video"),
            ("/uploads/<file>", "Serve uploaded CVs"),
        ],
    )

    # 12. Branding
    doc.add_heading("12. Branding & Colors", level=1)
    add_table(
        doc,
        ["Color", "Hex", "Usage"],
        [
            ("Teal", "#124e66", "Headers, primary brand"),
            ("Dark navy", "#0a2338", "Gradients, sidebar"),
            ("Orange", "#f39c12", "Buttons, accents, highlights"),
            ("Orange dark", "#e67e22", "Button hover"),
            ("Light gray", "#f5f7fa", "Page backgrounds"),
        ],
    )
    doc.add_paragraph("Main stylesheet: static/style.css")

    # 13. Troubleshooting
    doc.add_heading("13. Troubleshooting", level=1)
    add_table(
        doc,
        ["Problem", "Solution"],
        [
            ("App won't start", "Check Python version, run pip install -r requirements.txt"),
            ("Emails not sending", "Configure SMTP in .env or HR Settings; test from Settings"),
            ("Logo not showing", "Upload in HR Settings, or place file at static/gmb-logo.svg"),
            ("Welcome video not playing", "Use MP4 in uploads/branding/; upload via Settings"),
            ("Interviewer link expired", "HR must generate a new link (valid 24 hours)"),
            ("Wrong access code", "Use the code shown when HR generated the link"),
            ("Interview shows completed too early", "All panel interviewers must score; then HR screens"),
            ("Duplicate application", "Same email cannot apply twice to the same job"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("Support", level=1)
    add_bullets(doc, [
        "Company settings: HR → Settings in the dashboard",
        "Technical issues: Check data/activity_logs.json for recent actions",
        "Email problems: Verify .env SMTP credentials and test from Settings",
    ])

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Created: {path}")
