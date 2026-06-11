#!/usr/bin/env python
"""Generate ~50-page thesis Word document for GMB HR Recruitment System."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from thesis_content import (
    CHAPTER_1,
    CHAPTER_2,
    CHAPTER_3,
    CHAPTER_4,
    CHAPTER_5,
    ORGANIZATION,
    SUBTITLE,
    TITLE,
)
from thesis_expansions import EXPANSIONS, EXTRA_SECTIONS_CH3, EXTRA_SECTIONS_CH4, EXTRA_SECTIONS_CH5
from thesis_more_content import (
    CHAPTER_DISCUSSIONS,
    EXTENDED_APPENDICES,
    IMPLEMENTATION_DETAILS,
    COMPARATIVE_SYSTEMS,
    INSTITUTIONAL_CONTEXT,
    LIMITATIONS_AND_ETHICS,
    LITERATURE_REVIEW,
    METHODOLOGY_SECTION,
    TECHNICAL_DEEP_DIVE,
)

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "GMB_HR_SYSTEM_PROJECT_REPORT.docx"


def set_document_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(6)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(18, 78, 102)


def add_centered(doc, text, size=12, bold=False, color=None, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_paragraphs(doc, paragraphs, first_indent=True):
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if first_indent:
            p.paragraph_format.first_line_indent = Inches(0.5)


def merge_paragraphs(section_key, paragraphs):
    extra = EXPANSIONS.get(section_key, [])
    return paragraphs + extra


def add_section(doc, title, paragraphs, level=2):
    doc.add_heading(title, level=level)
    merged = merge_paragraphs(title, paragraphs)
    add_paragraphs(doc, merged)


def add_chapter(doc, chapter_num, title, sections, extra_sections=None):
    doc.add_heading(f"Chapter {chapter_num}: {title}", level=1)
    for section, paragraphs in sections.items():
        add_section(doc, section, paragraphs)
    if extra_sections:
        for section, paragraphs in extra_sections.items():
            add_section(doc, section, paragraphs)


def shade_cell(cell, fill="124E66"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        shade_cell(table.rows[0].cells[i])
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


ABSTRACT = [
    "This project report presents the development of an integrated web-based Human Resource Recruitment and Applicant Tracking System for the Grain Marketing Board (GMB). The study responds to operational challenges associated with manual recruitment processes, including fragmented application handling, inconsistent interview scoring, and limited auditability of hiring decisions. A comprehensive systems development methodology encompassing planning, analysis, design, and implementation was applied to deliver a multi-portal solution serving job applicants, HR personnel, external interview panel members, and audit stakeholders.",
    "The implemented system provides a public careers portal for vacancy publication and online applications, an HR dashboard for screening and interview management, an external interviewer portal with sequential candidate activation, and automated curriculum vitae scoring using text analysis techniques. Collaborative evaluation is supported through per-evaluator mark recording, average score computation, and rule-enforced screening that prevents final selection until all required interviewers have submitted assessments for all interviewed candidates.",
    "Testing and stakeholder review confirmed that the platform improves transparency, efficiency, and fairness in recruitment operations. The project contributes a practical reference model for public-sector HR technology adoption using Python Flask and modular service architecture. Recommendations for future work include enterprise database migration, enhanced analytics, and integration with broader human resource information systems.",
]

ACKNOWLEDGEMENTS = [
    "The successful completion of this project would not have been possible without the support and guidance of numerous individuals and institutions. The author expresses sincere gratitude to academic supervisors for their constructive feedback throughout the planning, analysis, design, and implementation phases. Special appreciation is extended to the Human Resources department of the Grain Marketing Board for providing domain expertise, requirements clarification, and participation in user acceptance testing.",
    "Thanks are also due to colleagues and peers who offered technical suggestions regarding system architecture, security practices, and user interface design. Finally, the author acknowledges the open-source software community whose frameworks and libraries—including Flask, scikit-learn, and python-docx—formed the technical foundation upon which this system was built.",
]

APPENDIX_A = [
    "Appendix A provides supplementary detail regarding functional requirements traceability. Each requirement identifier maps to design components and implementation modules to demonstrate completeness of coverage. FR-01 through FR-05 address applicant portal capabilities including vacancy listing, search, application capture, document upload, and duplicate prevention. FR-06 through FR-15 address HR job management, screening, interview scheduling, and communication. FR-16 through FR-25 address interviewer panel management, authentication, scoring, and sequential workflow control. FR-26 through FR-30 address audit logging, settings, branding, and reporting.",
    "Non-functional requirements NR-01 through NR-10 address performance, security, usability, maintainability, and portability as documented in Chapter Three. Verification methods included manual test scripts, inspection of route permission decorators, and review of JSON schema backfill migrations executed during application startup.",
]

APPENDIX_B = [
    "Appendix B summarises end-user procedures for key stakeholders. Applicants access the careers portal at the organisation's public URL, browse open positions, and complete the apply form with curriculum vitae attachment. HR managers log in through the HR System login page, navigate to the Jobs section to publish vacancies, and monitor incoming applications on the dashboard.",
    "Screening is initiated from the Screening section by selecting a job and executing run screening, which invokes CV scoring and updates candidate recommendations. Interview scheduling is performed from candidate rows with email template personalisation. External interviewers are registered under Interviews → External Interviewer Setup, after which HR generates a shared link and eight-digit access code distributed to all panel members.",
    "For each candidate, HR clicks Start to activate the interview; interviewers automatically see the active candidate on their portal and submit marks using the evaluation modal. After all candidates are interviewed and all panel members have scored, HR executes Screen Completed Interviews in the Interview Marks section and confirms selection of the highest average candidates.",
]

APPENDIX_C = [
    "Appendix C documents the principal source artefacts comprising the solution. The file app.py contains approximately three thousand lines implementing Flask routes, data persistence, enrichment utilities, permission checks, and recruitment business rules. The templates directory contains HTML pages for applicant, HR, interviewer, and audit interfaces. The services directory encapsulates curriculum vitae scoring, document validation, and email helper logic.",
    "Persistent data resides in the data directory as JSON files representing users, jobs, applications, settings, interviewer panels, activity logs, and email history. Uploaded applicant documents and branding media reside in the uploads directory. Static assets including style.css and default logo SVG files reside in the static directory. Configuration secrets are stored in environment variables rather than source control.",
]


def build_document():
    doc = Document()
    set_document_defaults(doc)

    # Title page
    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, ORGANIZATION, 14, bold=True, color=RGBColor(18, 78, 102), space_after=24)
    add_centered(doc, TITLE, 16, bold=True, space_after=18)
    add_centered(doc, SUBTITLE, 12, space_after=36)
    add_centered(doc, "June 2026", 12, space_after=6)
    doc.add_page_break()

    # Front matter
    doc.add_heading("Abstract", level=1)
    add_paragraphs(doc, ABSTRACT)
    doc.add_page_break()

    doc.add_heading("Acknowledgements", level=1)
    add_paragraphs(doc, ACKNOWLEDGEMENTS)
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc_lines = [
        "Chapter 1: Introduction",
        "Chapter 2: Planning",
        "Chapter 3: Analysis",
        "Chapter 4: Design",
        "Chapter 5: Implementation",
        "Appendices",
    ]
    for line in toc_lines:
        doc.add_paragraph(line)
    doc.add_page_break()

    doc.add_heading("List of Figures", level=1)
    figures = [
        "Figure 4.1: High-level three-tier system architecture diagram",
        "Figure 4.2: Applicant portal careers page wireframe",
        "Figure 4.3: HR dashboard navigation and sidebar layout",
        "Figure 4.4: Interview marks table showing multiple evaluators",
        "Figure 4.5: Sequential active candidate workflow diagram",
        "Figure 4.6: Interviewer portal active candidate view",
        "Figure 5.1: Screenshot of job posting form",
        "Figure 5.2: Screenshot of CV screening results",
        "Figure 5.3: Screenshot of interview scheduling modal",
        "Figure 5.4: Screenshot of screening ranking results",
    ]
    for fig in figures:
        doc.add_paragraph(fig)
    doc.add_paragraph(
        "Note: Insert actual screenshots from docs/images into the final submission copy of this document "
        "to accompany the figure list above. Screenshots substantially increase visual evidence and page depth."
    )
    doc.add_page_break()

    doc.add_heading("List of Tables", level=1)
    tables = [
        "Table 2.1: Project phase schedule and deliverables",
        "Table 3.1: Stakeholder goals and pain points",
        "Table 4.1: Module responsibility matrix",
        "Table 5.1: Test case results summary",
        "Table 5.2: Interview evaluation criteria rubric",
    ]
    for tbl in tables:
        doc.add_paragraph(tbl)
    doc.add_page_break()

    # Chapter 1
    ch1 = {**CHAPTER_1, **LITERATURE_REVIEW, **INSTITUTIONAL_CONTEXT, **LIMITATIONS_AND_ETHICS, "1.12 Discussion": CHAPTER_DISCUSSIONS["1.12 Discussion"]}
    add_chapter(doc, 1, "Introduction", ch1)
    doc.add_page_break()

    # Chapter 2
    ch2 = {**CHAPTER_2, **METHODOLOGY_SECTION, "2.11 Discussion of Planning Outcomes": CHAPTER_DISCUSSIONS["2.11 Discussion of Planning Outcomes"]}
    add_chapter(doc, 2, "Planning", ch2)
    add_table(
        doc,
        ["Phase", "Duration (Weeks)", "Key Deliverables"],
        [
            ("Initiation & Planning", "2", "Proposal, scope, stakeholder map"),
            ("Systems Analysis", "3", "Requirements specification, process models"),
            ("System Design", "3", "Architecture, UI mock-ups, module specs"),
            ("Implementation", "5", "Working application, services, templates"),
            ("Testing & Documentation", "2", "Test results, user guides, project report"),
        ],
    )
    doc.add_page_break()

    # Chapter 3
    ch3 = {**CHAPTER_3, **EXTRA_SECTIONS_CH3, **COMPARATIVE_SYSTEMS, "3.11 Discussion of Analytical Findings": CHAPTER_DISCUSSIONS["3.11 Discussion of Analytical Findings"]}
    add_chapter(doc, 3, "Analysis", ch3)
    add_table(
        doc,
        ["Actor", "Primary Goals", "Pain Points Addressed"],
        [
            ("Job Applicant", "Find vacancies and apply online", "Unclear process, repeated submissions"),
            ("HR Manager", "Hire qualified staff fairly", "Manual collation, weak audit trail"),
            ("HR Staff", "Screen and schedule interviews", "Lost CVs, email overload"),
            ("External Interviewer", "Score candidates conveniently", "Late access, informal mark submission"),
            ("Audit Officer", "Verify process integrity", "Incomplete records"),
        ],
    )
    doc.add_page_break()

    # Chapter 4
    ch4 = {**CHAPTER_4, **EXTRA_SECTIONS_CH4, "4.12 Discussion of Design Decisions": CHAPTER_DISCUSSIONS["4.12 Discussion of Design Decisions"]}
    add_chapter(doc, 4, "Design", ch4)
    add_table(
        doc,
        ["Module", "Primary Files", "Responsibility"],
        [
            ("Authentication", "app.py, login.html", "Login, sessions, roles"),
            ("Applications", "apply.html, app.py", "Capture and store applications"),
            ("Screening", "services/cv_scoring.py", "CV parsing and ranking"),
            ("Interviews", "hr_dashboard.html, app.py", "Schedule, start, score, screen"),
            ("Interviewer Panel", "interviewer_*.html, app.py", "External panel access"),
            ("Settings", "hr_settings.html, app.py", "Branding and SMTP"),
        ],
    )
    doc.add_page_break()

    # Chapter 5
    ch5 = {**CHAPTER_5, **IMPLEMENTATION_DETAILS, **EXTRA_SECTIONS_CH5, **TECHNICAL_DEEP_DIVE, "5.12 Discussion of Implementation Experience": CHAPTER_DISCUSSIONS["5.12 Discussion of Implementation Experience"]}
    add_chapter(doc, 5, "Implementation", ch5)
    add_table(
        doc,
        ["Test Case", "Expected Result", "Outcome"],
        [
            ("Applicant applies twice same email/job", "Second application rejected", "Pass"),
            ("Interviewer uses wrong access code", "Login denied", "Pass"),
            ("Single evaluator marks candidate", "Status not completed", "Pass"),
            ("All evaluators mark all candidates", "Screening permitted", "Pass"),
            ("Screening executed", "Ranked by average score", "Pass"),
            ("Select best candidates", "Top averages selected", "Pass"),
        ],
    )
    doc.add_page_break()

    # Appendices
    doc.add_heading("Appendix A: Requirements Traceability", level=1)
    add_paragraphs(doc, APPENDIX_A)
    doc.add_page_break()

    doc.add_heading("Appendix B: User Procedures Summary", level=1)
    add_paragraphs(doc, APPENDIX_B)
    doc.add_page_break()

    doc.add_heading("Appendix C: Source Code and Directory Structure", level=1)
    add_paragraphs(doc, APPENDIX_C)
    doc.add_page_break()

    for title, paragraphs in EXTENDED_APPENDICES.items():
        doc.add_heading(title, level=1)
        add_paragraphs(doc, paragraphs)
        doc.add_page_break()

    doc.add_heading("References", level=1)
    references = [
        "Flask Documentation (2026). Pallets Projects. https://flask.palletsprojects.com/",
        "Sommerville, I. (2016). Software Engineering. Pearson Education.",
        "Dennis, A., Wixom, B., & Roth, R. (2014). Systems Analysis and Design. Wiley.",
        "Scikit-learn Developers (2026). TF-IDF Vectorization Documentation.",
        "Werkzeug Security Documentation (2026). Password Hashing Utilities.",
        "Grain Marketing Board Internal HR Policy Documents (consultation materials).",
    ]
    for ref in references:
        doc.add_paragraph(ref)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    # Rough word count estimate
    from docx import Document as D

    d = D(path)
    words = sum(len(p.text.split()) for p in d.paragraphs)
    print(f"Created: {path}")
    print(f"Estimated words: {words}")
    print(f"Estimated pages (~300 words/page): {words / 300:.1f}")
